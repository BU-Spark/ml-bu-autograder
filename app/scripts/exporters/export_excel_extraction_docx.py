#!/usr/bin/env python3
"""
Create readable DOCX files from per-file extraction JSON (XLSX sources only).

For each *.xlsx.extraction.json:
- Reconstruct table chunks by sheet and row range.
- Write a structured Word table for easy human review.
- Include extracted text/summary for any embedded image chunks.
"""

from __future__ import annotations

import argparse
import json
import re
from pathlib import Path
from typing import Any

try:
    from docx import Document
except Exception as exc:  # pragma: no cover
    raise SystemExit(
        "python-docx is required. Install with: "
        "pip install python-docx (or pip install -r requirements.txt)"
    ) from exc


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(description="Export Excel extraction JSON to DOCX.")
    parser.add_argument(
        "--per-file-json-dir",
        required=True,
        help="Root folder containing *.extraction.json files from pipeline output.",
    )
    parser.add_argument(
        "--output-dir",
        required=True,
        help="Folder where DOCX exports should be written.",
    )
    return parser.parse_args()


def safe_mkdir(path: Path) -> None:
    path.mkdir(parents=True, exist_ok=True)


def clean_text(text: str) -> str:
    text = text.replace("\x00", " ")
    text = text.replace("\u2014", "-")
    text = re.sub(r"\s+", " ", text).strip()
    return text


def parse_table_chunk(content: str) -> tuple[str, list[str], list[dict[str, str]]]:
    """
    Parse pipeline table chunk format:
    Sheet: <name>
    Headers: A | B | C
    Rows:
    row_18: A: x ; C: y
    """
    lines = [line.rstrip() for line in str(content).splitlines()]
    sheet_name = ""
    headers: list[str] = []
    rows: list[dict[str, str]] = []

    in_rows = False
    for line in lines:
        s = line.strip()
        if not s:
            continue
        if s.startswith("Sheet:"):
            sheet_name = clean_text(s.split(":", 1)[1])
            continue
        if s.startswith("Headers:"):
            hdr = s.split(":", 1)[1]
            headers = [clean_text(h) for h in hdr.split("|")]
            continue
        if s == "Rows:":
            in_rows = True
            continue
        if in_rows and s.startswith("row_"):
            # row_18: key: val ; key2: val2
            _, payload = s.split(":", 1)
            entry = {"__row_id": clean_text(s.split(":", 1)[0])}
            parts = [p.strip() for p in payload.split(";") if p.strip()]
            for p in parts:
                if ":" in p:
                    k, v = p.split(":", 1)
                    entry[clean_text(k)] = clean_text(v)
            rows.append(entry)

    return sheet_name, headers, rows


def image_text_from_chunk(chunk: dict[str, Any]) -> str:
    md = chunk.get("metadata", {})
    extracted = clean_text(str(md.get("extracted_text", "") or ""))
    if extracted:
        return extracted
    desc = clean_text(str(md.get("description", "") or ""))
    if desc:
        return desc
    return clean_text(str(chunk.get("content", "") or "")) or "[no_image_text_extracted]"


def chunk_sort_key(chunk: dict[str, Any]) -> tuple[str, int]:
    md = chunk.get("metadata", {})
    sheet = str(md.get("sheet_name") or "")
    start = int(md.get("row_start") or md.get("image_index_on_sheet") or 0)
    return (sheet, start)


def export_excel_json_to_docx(json_path: Path, per_file_json_root: Path, output_dir: Path) -> Path | None:
    payload = json.loads(json_path.read_text(encoding="utf-8"))
    source_path = str(payload.get("source_path", ""))
    if not source_path.lower().endswith(".xlsx"):
        return None

    chunks = payload.get("chunks", [])
    if not isinstance(chunks, list):
        chunks = []
    chunks = sorted(chunks, key=chunk_sort_key)

    rel = json_path.relative_to(per_file_json_root)
    stem = str(rel)
    if stem.endswith(".extraction.json"):
        stem = stem[: -len(".extraction.json")]
    out_path = output_dir / f"{stem}.docx"
    safe_mkdir(out_path.parent)

    doc = Document()
    doc.add_heading(f"Excel Extraction Reconstruction: {Path(source_path).name}", level=1)
    doc.add_paragraph(f"Source: {source_path}")

    current_sheet = None
    image_counter = 0
    for chunk in chunks:
        md = chunk.get("metadata", {})
        ctype = md.get("content_type", "")
        sheet = str(md.get("sheet_name") or "Unknown Sheet")
        if sheet != current_sheet:
            current_sheet = sheet
            doc.add_heading(f"Sheet: {sheet}", level=2)

        if ctype == "table":
            title = f"Table rows {md.get('row_start', '?')} to {md.get('row_end', '?')}"
            doc.add_paragraph(title)
            _, headers, rows = parse_table_chunk(str(chunk.get("content", "") or ""))
            if not headers:
                # Fallback when parsing misses format.
                doc.add_paragraph(clean_text(str(chunk.get("content", "") or "")))
                continue

            # Include row id as first column so row boundaries are visible.
            cols = ["__row_id"] + headers
            table = doc.add_table(rows=1, cols=len(cols))
            table.style = "Table Grid"
            for i, h in enumerate(cols):
                table.cell(0, i).text = h
            for row in rows:
                cells = table.add_row().cells
                for i, h in enumerate(cols):
                    cells[i].text = str(row.get(h, ""))

        elif ctype == "image_description":
            image_counter += 1
            txt = image_text_from_chunk(chunk)
            p = doc.add_paragraph()
            p.add_run(f"[Embedded Image {image_counter} extracted text]").bold = True
            p.add_run("\n" + txt)

    doc.save(out_path)
    return out_path


def main() -> int:
    args = parse_args()
    per_file_json_dir = Path(args.per_file_json_dir).expanduser().resolve()
    output_dir = Path(args.output_dir).expanduser().resolve()

    if not per_file_json_dir.exists() or not per_file_json_dir.is_dir():
        raise SystemExit(f"Invalid per-file-json-dir: {per_file_json_dir}")

    safe_mkdir(output_dir)

    files = sorted(per_file_json_dir.rglob("*.extraction.json"))
    written: list[Path] = []
    for fp in files:
        out = export_excel_json_to_docx(fp, per_file_json_dir, output_dir)
        if out is not None:
            written.append(out)

    print(f"Input extraction files found: {len(files)}")
    print(f"Excel DOCX files generated: {len(written)}")
    if written:
        print("Sample outputs:")
        for p in written[:5]:
            print(str(p))
    return 0


if __name__ == "__main__":
    raise SystemExit(main())

