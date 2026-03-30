#!/usr/bin/env python3
"""
Export PDF extraction JSON to DOCX with layout-style reconstruction.

For each *.pdf.extraction.json:
- Keep extracted text in document order.
- Insert the extracted image at the corresponding image position.
- Add optional extracted image text and model description beneath the image.
"""

from __future__ import annotations

import argparse
import json
import re
from pathlib import Path
from typing import Any

try:
    from docx import Document
    from docx.shared import Inches
except Exception as exc:  # pragma: no cover
    raise SystemExit(
        "python-docx is required. Install with: "
        "pip install python-docx (or pip install -r requirements.txt)"
    ) from exc


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(
        description="Export PDF extraction JSON to DOCX with images in-place."
    )
    parser.add_argument(
        "--per-file-json-dir",
        required=True,
        help="Root folder containing *.extraction.json files from pipeline output.",
    )
    parser.add_argument(
        "--output-dir",
        required=True,
        help="Folder where DOCX exports should be written.",
    )
    parser.add_argument(
        "--include-image-text",
        action="store_true",
        help="Deprecated: image text is always included.",
    )
    parser.add_argument(
        "--include-image-summary",
        action="store_true",
        help="Include model image summary below each inserted image.",
    )
    parser.add_argument(
        "--image-width-inches",
        type=float,
        default=6.2,
        help="Target width for inserted images in inches (default: 6.2).",
    )
    return parser.parse_args()


def safe_mkdir(path: Path) -> None:
    path.mkdir(parents=True, exist_ok=True)


def clean_text(text: str) -> str:
    text = text.replace("\x00", " ")
    text = text.replace("\u2014", "-")
    text = re.sub(r"\s+", " ", text).strip()
    return text


def chunk_sort_key(chunk: dict[str, Any]) -> tuple[int, int, int]:
    md = chunk.get("metadata", {})
    doc_order = md.get("document_order")
    page = md.get("page_number") or md.get("page") or 0
    block_index = md.get("block_index") or 0

    if isinstance(doc_order, int) and doc_order > 0:
        return (0, doc_order, 0)
    return (1, int(page), int(block_index))


def image_text_from_chunk(chunk: dict[str, Any]) -> str:
    md = chunk.get("metadata", {})
    extracted = clean_text(str(md.get("extracted_text", "") or ""))
    if extracted:
        return extracted

    visible = md.get("visible_text_lines")
    if isinstance(visible, list):
        lines = [clean_text(str(x)) for x in visible if clean_text(str(x))]
        if lines:
            return "\n".join(lines)

    return ""


def insert_image(doc: Document, image_path: str, width_in_inches: float) -> bool:
    p = Path(image_path)
    if not p.exists() or not p.is_file():
        return False
    try:
        doc.add_picture(str(p), width=Inches(width_in_inches))
        return True
    except Exception:
        return False


def export_pdf_json_to_docx(
    json_path: Path,
    per_file_json_root: Path,
    output_dir: Path,
    include_image_text: bool,
    include_image_summary: bool,
    image_width_in_inches: float,
) -> Path | None:
    payload = json.loads(json_path.read_text(encoding="utf-8"))
    source_path = str(payload.get("source_path", ""))
    if not source_path.lower().endswith(".pdf"):
        return None

    chunks = payload.get("chunks", [])
    if not isinstance(chunks, list):
        chunks = []
    chunks = sorted(chunks, key=chunk_sort_key)

    rel = json_path.relative_to(per_file_json_root)
    stem = str(rel)
    if stem.endswith(".extraction.json"):
        stem = stem[: -len(".extraction.json")]
    out_path = output_dir / f"{stem}.docx"
    safe_mkdir(out_path.parent)

    doc = Document()
    doc.add_heading(f"Reconstructed Extraction: {Path(source_path).name}", level=1)
    doc.add_paragraph(f"Source: {source_path}")

    current_page = None
    image_counter = 0

    for chunk in chunks:
        md = chunk.get("metadata", {})
        page = md.get("page_number") or md.get("page")
        if isinstance(page, int) and page != current_page:
            current_page = page
            doc.add_heading(f"Page {page}", level=2)

        ctype = md.get("content_type", "")
        if ctype == "text":
            text = clean_text(str(chunk.get("content", "") or ""))
            if text:
                doc.add_paragraph(text)
            continue

        if ctype != "image_description":
            continue

        image_counter += 1
        image_path = str(md.get("image_path", "") or "")

        header = doc.add_paragraph()
        header.add_run(f"[Image {image_counter}]").bold = True

        inserted = insert_image(doc, image_path, image_width_in_inches)
        if not inserted:
            doc.add_paragraph(f"[image not found] {image_path}")

        if include_image_text:
            img_text = image_text_from_chunk(chunk)
            if img_text:
                p = doc.add_paragraph()
                p.add_run("Extracted image text:").bold = True
                p.add_run("\n" + img_text)

        if include_image_summary:
            summary = clean_text(str(md.get("description", "") or ""))
            if summary:
                p = doc.add_paragraph()
                p.add_run("Image summary:").bold = True
                p.add_run("\n" + summary)

    doc.save(out_path)
    return out_path


def main() -> int:
    args = parse_args()
    per_file_json_dir = Path(args.per_file_json_dir).expanduser().resolve()
    output_dir = Path(args.output_dir).expanduser().resolve()

    if not per_file_json_dir.exists() or not per_file_json_dir.is_dir():
        raise SystemExit(f"Invalid per-file-json-dir: {per_file_json_dir}")

    safe_mkdir(output_dir)

    files = sorted(per_file_json_dir.rglob("*.extraction.json"))
    written: list[Path] = []
    for fp in files:
        out = export_pdf_json_to_docx(
            json_path=fp,
            per_file_json_root=per_file_json_dir,
            output_dir=output_dir,
            include_image_text=True,
            include_image_summary=args.include_image_summary,
            image_width_in_inches=args.image_width_inches,
        )
        if out is not None:
            written.append(out)

    print(f"Input extraction files found: {len(files)}")
    print(f"PDF layout DOCX files generated: {len(written)}")
    if written:
        print("Sample outputs:")
        for p in written[:5]:
            print(str(p))
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
