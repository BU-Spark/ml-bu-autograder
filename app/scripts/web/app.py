#!/usr/bin/env python3
"""
app.py  —  Flask web application for the AI Auto Grader.

Run:
    cd "Final AI Auto Grader"
    python scripts/web/app.py

Then open:  http://localhost:5000
"""
from __future__ import annotations

import csv
import io
import json
import os
import re
import subprocess
import sys
import uuid
from datetime import datetime
from pathlib import Path

from flask import Flask, render_template, request, jsonify, Response, send_file
from dotenv import load_dotenv
from werkzeug.utils import secure_filename

# ── project paths ──────────────────────────────────────────────────────────
PROJECT_ROOT = Path(__file__).resolve().parents[2]
SCRIPTS_DIR = PROJECT_ROOT / "scripts"
RUN_PIPELINE = SCRIPTS_DIR / "cli" / "run_pipeline.py"
OUTPUT_ROOT = PROJECT_ROOT / "outputs" / "final_phase1"
DEFAULT_LECTURE_CHUNKS = (
    # Hybrid lecture chunks: HTML text blocks + PDF image descriptions (best of both)
    # HTML gives clean structured text; PDF image chunks add 91 AI-described diagrams/figures
    # that HTML completely misses (all img alt tags are empty, image files not on disk).
    OUTPUT_ROOT / "lecture_chunks_hybrid.jsonl"
)
DEFAULT_ASSIGNMENT = PROJECT_ROOT / "assignments" / "assignment1_instructions.txt"
DEFAULT_RUBRIC_DIR = Path(
    os.getenv("AUTO_GRADER_RUBRIC_DIR", "/Users/sai/Downloads/Spring 2026 2/Assignment Rubrics")
).expanduser()

# Library dirs — professor uploads here once; files persist across sessions
LIBRARY_DIR             = PROJECT_ROOT / "data" / "library"
LIBRARY_ASSIGNMENTS_DIR = LIBRARY_DIR / "assignments"
LIBRARY_QUIZZES_DIR     = LIBRARY_DIR / "quizzes"
LIBRARY_RUBRICS_DIR     = LIBRARY_DIR / "rubrics"

# Shared Chroma DB for lecture content — built once, reused across all grading runs.
# Separate DB paths per embedding provider so switching never corrupts an existing index.
def _embedding_provider() -> str:
    """
    Resolve which embedding provider to use.
    Priority: explicit CHROMA_EMBEDDING_PROVIDER env var → Google (if key present) → default.
    """
    explicit = os.getenv("CHROMA_EMBEDDING_PROVIDER", "").strip().lower()
    if explicit in {"openai", "google", "default"}:
        return explicit
    # Auto-select Google when a Gemini/Google API key is available — free and higher quality
    # than the default local embeddings.
    if os.getenv("GEMINI_API_KEY") or os.getenv("GOOGLE_API_KEY"):
        return "google"
    return "default"

def _shared_chroma_dir() -> Path:
    return OUTPUT_ROOT / f"shared_lecture_chroma_{_embedding_provider()}"

def _embedding_env() -> dict[str, str]:
    """Pass the resolved provider to child pipeline processes via env var."""
    return {"CHROMA_EMBEDDING_PROVIDER": _embedding_provider()}


def _adaptive_top_k() -> int:
    """Return retrieval top-k scaled to lecture corpus size.
    Small corpus (≤100 chunks) → 4, medium (≤500) → 6, large (>500) → 8.
    This avoids flooding context with noise on small corpora or missing
    relevant chunks on large ones.
    """
    try:
        chunks_path = DEFAULT_LECTURE_CHUNKS
        if not chunks_path.exists():
            return 6
        count = sum(1 for line in chunks_path.read_text(encoding="utf-8").splitlines() if line.strip())
        if count <= 100:
            return 4
        if count <= 500:
            return 6
        return 8
    except Exception:
        return 6

SHARED_LECTURE_CHROMA = OUTPUT_ROOT / "shared_lecture_chroma"  # legacy (unused)
SHARED_LECTURE_COLLECTION = "lecture_v1"

STUDENT_ALLOWED_EXTS = {".pdf", ".pptx", ".xlsx"}
SUPPORT_ALLOWED_EXTS     = {".docx", ".pdf", ".txt", ".md"}
RUBRIC_ALLOWED_EXTS      = SUPPORT_ALLOWED_EXTS | {".json"}  # generated rubrics are JSON

# Load environment from project .env so web app can access API keys without manual exports.
load_dotenv(PROJECT_ROOT / ".env", override=True)

# ── provider config ────────────────────────────────────────────────────────
PROVIDERS = {
    "openai": {
        "label": "OpenAI GPT-4o",
        "model": "gpt-4o-mini",
        "color": "#10a37f",
        "icon": "openai",
    },
    "gemini": {
        "label": "Google Gemini",
        "model": "gemini-2.5-flash",
        "color": "#4285f4",
        "icon": "gemini",
    },
    "anthropic": {
        "label": "Anthropic Claude",
        "model": "claude-sonnet-4-6",
        "color": "#7c3aed",
        "icon": "anthropic",
    },
}

PROVIDER_API_KEY_ENV = {
    "openai": "OPENAI_API_KEY",
    "gemini": "GEMINI_API_KEY",
    "anthropic": "ANTHROPIC_API_KEY",
}


# ── Flask app ──────────────────────────────────────────────────────────────
app = Flask(__name__)
app.config["MAX_CONTENT_LENGTH"] = 200 * 1024 * 1024  # 200 MB max


def _safe(name: str) -> str:
    return re.sub(r"[^a-zA-Z0-9._-]+", "_", name)


def _is_allowed_ext(filename: str, allowed_exts: set[str]) -> bool:
    ext = Path(filename or "").suffix.lower()
    return ext in allowed_exts


def _safe_upload_name(filename: str) -> str:
    base = secure_filename(filename or "") or "upload.bin"
    return f"{uuid.uuid4().hex[:8]}_{base}"


def _resolve_selected_path(selected: str, *, allowed_roots: list[Path], allowed_exts: set[str]) -> Path | None:
    if not selected:
        return None
    try:
        p = Path(selected).expanduser().resolve()
    except Exception:
        return None
    if not p.exists() or not p.is_file():
        return None
    if p.suffix.lower() not in allowed_exts:
        return None
    for root in allowed_roots:
        try:
            root_resolved = root.expanduser().resolve()
            p.relative_to(root_resolved)
            return p
        except Exception:
            continue
    return None


def _resolve_project_file(selected: str, *, allowed_exts: set[str]) -> Path | None:
    if not selected:
        return None
    try:
        p = Path(selected).expanduser().resolve()
    except Exception:
        return None
    if not p.exists() or not p.is_file():
        return None
    if p.suffix.lower() not in allowed_exts:
        return None
    try:
        p.relative_to(PROJECT_ROOT.resolve())
    except Exception:
        return None
    return p


def _run(args: list[str], extra_env: dict | None = None) -> tuple[int, str]:
    env = os.environ.copy()
    if extra_env:
        env.update(extra_env)
    try:
        proc = subprocess.run(
            args,
            cwd=str(PROJECT_ROOT),
            env=env,
            text=True,
            capture_output=True,
            timeout=600,
        )
    except subprocess.TimeoutExpired:
        return 1, "Pipeline step timed out after 10 minutes."
    return proc.returncode, ((proc.stdout or "") + "\n" + (proc.stderr or "")).strip()


def _cli() -> list[str]:
    return [sys.executable, str(RUN_PIPELINE)]


def _collect_supporting_files(run_root: Path) -> tuple[Path | None, Path | None, str | None]:
    rubric_file = request.files.get("rubric")
    assignment_file = request.files.get("assignment")
    selected_rubric = request.form.get("selected_rubric", "")
    selected_assignment = request.form.get("selected_assignment", "")

    if rubric_file and rubric_file.filename and not _is_allowed_ext(rubric_file.filename, RUBRIC_ALLOWED_EXTS):
        return None, None, "Invalid rubric file type. Allowed: DOCX, PDF, TXT, MD, JSON."
    if assignment_file and assignment_file.filename and not _is_allowed_ext(assignment_file.filename, SUPPORT_ALLOWED_EXTS):
        return None, None, "Invalid assignment file type. Allowed: DOCX, PDF, TXT, MD."

    support_dir = run_root / "supporting"
    rubric_path = None
    assignment_path = None

    generated_rubric_json = request.form.get("generated_rubric_json", "").strip()
    if generated_rubric_json:
        try:
            json.loads(generated_rubric_json)
        except Exception:
            return None, None, "generated_rubric_json is not valid JSON."
        support_dir.mkdir(parents=True, exist_ok=True)
        rubric_path = support_dir / "generated_rubric.json"
        rubric_path.write_text(generated_rubric_json, encoding="utf-8")
    elif rubric_file and rubric_file.filename:
        support_dir.mkdir(parents=True, exist_ok=True)
        rubric_path = support_dir / _safe_upload_name(rubric_file.filename)
        rubric_file.save(str(rubric_path))
    elif selected_rubric:
        rubric_path = _resolve_selected_path(
            selected_rubric,
            allowed_roots=[DEFAULT_RUBRIC_DIR, PROJECT_ROOT / "assignments",
                           LIBRARY_RUBRICS_DIR, LIBRARY_ASSIGNMENTS_DIR, LIBRARY_QUIZZES_DIR],
            allowed_exts=RUBRIC_ALLOWED_EXTS,
        )

    if assignment_file and assignment_file.filename:
        support_dir.mkdir(parents=True, exist_ok=True)
        assignment_path = support_dir / _safe_upload_name(assignment_file.filename)
        assignment_file.save(str(assignment_path))
    elif selected_assignment:
        assignment_path = _resolve_selected_path(
            selected_assignment,
            allowed_roots=[PROJECT_ROOT / "assignments", DEFAULT_RUBRIC_DIR,
                           LIBRARY_ASSIGNMENTS_DIR, LIBRARY_QUIZZES_DIR, LIBRARY_RUBRICS_DIR],
            allowed_exts=SUPPORT_ALLOWED_EXTS,
        )
    # No fallback to DEFAULT_ASSIGNMENT — professor must explicitly choose

    # ── Strict validation: both required ──────────────────────────
    if not rubric_path:
        return None, None, (
            "No rubric selected. Please choose a rubric from the library "
            "or upload one before grading."
        )
    if not assignment_path:
        return None, None, (
            "No assignment description selected. Please choose an assignment "
            "from the library or upload one before grading."
        )

    return rubric_path, assignment_path, None


# ── Discover available rubric / assignment files ───────────────────────────
def _display_name(filename: str, kind: str = "description") -> str:
    """Turn a filename into a human-readable label.
    assignment_1.pdf        → Assignment - 1 - Description
    assignment_1_rubric.json → Assignment - 1 - Rubric
    quiz_1.pdf              → Quiz - 1 - Description
    quiz_1_rubric.json      → Quiz - 1 - Rubric
    """
    stem = Path(filename).stem
    is_rubric = bool(re.search(r"(?i)rubric", stem))
    label = "Rubric" if (is_rubric or kind == "rubric") else "Description"
    m_a = re.match(r"(?i)assignment[_\-\s]*(\d+)", stem)
    if m_a:
        return f"Assignment - {m_a.group(1)} - {label}"
    m_q = re.match(r"(?i)quiz[_\-\s]*(\d+)", stem)
    if m_q:
        return f"Quiz - {m_q.group(1)} - {label}"
    # Fallback: clean up underscores/dashes
    name = re.sub(r"[-_]+", " ", stem).strip()
    return " ".join(w.capitalize() for w in name.split())


def _assignment_number(filename: str) -> str | None:
    """Extract number from assignment_N filename. assignment_1.pdf → '1'"""
    m = re.match(r"(?i)assignment[_\-\s]*(\d+)", Path(filename).stem)
    return m.group(1) if m else None


def _quiz_number(filename: str) -> str | None:
    """Extract number from quiz_N filename. quiz_1.pdf → '1'"""
    m = re.match(r"(?i)quiz[_\-\s]*(\d+)", Path(filename).stem)
    return m.group(1) if m else None


def _discover_files(root: Path, exts: set[str] | None = None, kind: str = "description") -> list[dict]:
    if not root.exists():
        return []
    exts = exts or {".pdf", ".docx", ".txt", ".md"}
    out: list[dict] = []
    for p in sorted(root.rglob("*")):
        if p.is_file() and p.suffix.lower() in exts and not p.name.startswith("~$"):
            out.append({
                "path": str(p),
                "name": p.name,
                "rel": str(p.relative_to(root)),
                "display_name": _display_name(p.name, kind=kind),
                "assignment_number": _assignment_number(p.name),
                "quiz_number": _quiz_number(p.name),
            })
    return out


def _shared_chroma_ready() -> bool:
    """True if the shared lecture Chroma DB has been built and contains data."""
    chroma_dir = _shared_chroma_dir()
    chroma_sqlite = chroma_dir / "chroma.sqlite3"
    return chroma_dir.exists() and chroma_sqlite.exists()


# ── API routes ─────────────────────────────────────────────────────────────
@app.route("/")
def index():
    rubrics = (
        _discover_files(DEFAULT_RUBRIC_DIR, kind="rubric") +
        _discover_files(LIBRARY_RUBRICS_DIR, exts=RUBRIC_ALLOWED_EXTS, kind="rubric")
    )
    assignments = (
        _discover_files(PROJECT_ROOT / "assignments") +
        _discover_files(LIBRARY_ASSIGNMENTS_DIR)
    )
    quizzes = _discover_files(LIBRARY_QUIZZES_DIR)
    return render_template(
        "index.html",
        providers=PROVIDERS,
        rubrics=rubrics,
        assignments=assignments,
        quizzes=quizzes,
        lecture_chunks_exist=DEFAULT_LECTURE_CHUNKS.exists(),
        shared_chroma_ready=_shared_chroma_ready(),
    )


@app.route("/api/providers")
def api_providers():
    return jsonify(PROVIDERS)


@app.route("/api/index-lectures", methods=["POST"])
def api_index_lectures():
    """Pre-index lecture chunks into the shared Chroma DB.
    Call this once before grading any students. Subsequent grading runs
    reuse this DB instead of rebuilding it from scratch every time.
    """
    if not DEFAULT_LECTURE_CHUNKS.exists():
        return jsonify(success=False, error=f"Lecture chunks not found at: {DEFAULT_LECTURE_CHUNKS}"), 404

    chroma_dir = _shared_chroma_dir()
    chroma_dir.mkdir(parents=True, exist_ok=True)
    code, out = _run(
        _cli() + [
            "--mode", "index",
            "--chunks-jsonl", str(DEFAULT_LECTURE_CHUNKS),
            "--chroma-path", str(chroma_dir),
            "--chroma-collection", SHARED_LECTURE_COLLECTION,
            "--output-root", str(OUTPUT_ROOT),
            "--run-id", "shared_lecture_index",
        ],
        extra_env=_embedding_env(),
    )
    if code != 0:
        return jsonify(success=False, error="Lecture indexing failed.", log=out[-3000:]), 500
    return jsonify(success=True, message="Lecture index built successfully.", log=out[-1000:])


@app.errorhandler(413)
def payload_too_large(_err):
    return jsonify(success=False, error="Uploaded file is too large (max 200 MB)."), 413


@app.route("/api/grade", methods=["POST"])
def api_grade():
    student = (
        request.files.get("student_pdf")
        or request.files.get("student_submission")
        or request.files.get("student_file")
    )
    if not student or not student.filename:
        return jsonify(success=False, error="Please upload a student submission file."), 400
    if not _is_allowed_ext(student.filename, STUDENT_ALLOWED_EXTS):
        return jsonify(success=False, error="Invalid student file type. Allowed: PDF, PPTX, XLSX."), 400

    describe_provider = request.form.get("describe_provider", "openai")
    describe_model = request.form.get("describe_model") or PROVIDERS.get(
        describe_provider, {}
    ).get("model", "gpt-4o-mini")
    grade_provider = request.form.get("grade_provider", "openai")
    grade_model = request.form.get("grade_model") or PROVIDERS.get(
        grade_provider, {}
    ).get("model", "gpt-4o-mini")

    for provider_name in {describe_provider, grade_provider}:
        key_env_name = PROVIDER_API_KEY_ENV.get(provider_name)
        if key_env_name and not os.getenv(key_env_name):
            return jsonify(
                success=False,
                error=f"{key_env_name} is not set. Add it to the project .env file or export it before starting Flask.",
            ), 400

    # ── create run directory ──
    run_id = f"web_{datetime.now().strftime('%Y%m%d_%H%M%S')}_{uuid.uuid4().hex[:4]}"
    run_root = OUTPUT_ROOT / run_id
    upload_dir = run_root / "uploads"
    upload_dir.mkdir(parents=True, exist_ok=True)

    original_student_name = student.filename   # preserve before UUID-prefix rename
    student_path = upload_dir / _safe_upload_name(student.filename)
    student.save(str(student_path))

    rubric_path, assignment_path, support_error = _collect_supporting_files(run_root)
    if support_error:
        return jsonify(success=False, error=support_error), 400

    safe_describe_model = _safe(describe_model)
    extract_dir = run_root / "extract"
    describe_dir = run_root / f"describe_student_{describe_provider}_{safe_describe_model}"
    retrieval_out = run_root / "retrieval.jsonl"

    # Use shared Chroma DB if it exists; otherwise fall back to per-run DB
    use_shared_chroma = _shared_chroma_ready() and DEFAULT_LECTURE_CHUNKS.exists()
    chroma_path = _shared_chroma_dir() if use_shared_chroma else run_root / "chroma_db"

    steps: list[dict] = []

    # ── 1. Extract ──
    code, out = _run(
        _cli() + [
            "--mode", "extract",
            "--data-dir", str(upload_dir),
            "--output-root", str(OUTPUT_ROOT),
            "--run-id", run_id,
        ]
    )
    steps.append({"step": "Extract PDF", "ok": code == 0, "log": out[-2000:]})
    if code != 0:
        return jsonify(success=False, error="PDF extraction failed.", steps=steps, log=out[-1000:])

    # ── 2. Describe ──
    code, out = _run(
        _cli() + [
            "--mode", "describe",
            "--extract-dir", str(extract_dir),
            "--describe-dir", str(describe_dir),
            "--vision-provider", describe_provider,
            "--vision-model", describe_model,
            "--prompt-version", "verbose_v2",
        ]
    )
    steps.append({"step": "Analyze Content", "ok": code == 0, "log": out[-2000:]})
    if code != 0:
        return jsonify(success=False, error="Content analysis failed.", steps=steps, log=out[-1000:])

    # ── 3. Index lectures (only if shared DB not ready) ──
    has_lectures = DEFAULT_LECTURE_CHUNKS.exists()
    if has_lectures:
        if not use_shared_chroma:
            # No shared DB yet — build it now (also saves it as shared for future runs)
            chroma_dir = _shared_chroma_dir()
            chroma_dir.mkdir(parents=True, exist_ok=True)
            code, out = _run(
                _cli() + [
                    "--mode", "index",
                    "--chunks-jsonl", str(DEFAULT_LECTURE_CHUNKS),
                    "--output-root", str(OUTPUT_ROOT),
                    "--run-id", run_id,
                    "--chroma-path", str(chroma_dir),
                    "--chroma-collection", SHARED_LECTURE_COLLECTION,
                ],
                extra_env=_embedding_env(),
            )
            steps.append({"step": "Index Lectures", "ok": code == 0, "log": out[-2000:]})
            if code != 0:
                return jsonify(success=False, error="Lecture indexing failed.", steps=steps)
        else:
            steps.append({"step": "Index Lectures", "ok": True, "log": "Reusing shared lecture index (skipped rebuild)."})

        # ── 4. Retrieve ──
        code, out = _run(
            _cli() + [
                "--mode", "retrieve",
                "--chunks-jsonl", str(describe_dir / "chunks.jsonl"),
                "--output-root", str(OUTPUT_ROOT),
                "--run-id", run_id,
                "--chroma-path", str(chroma_path),
                "--chroma-collection", SHARED_LECTURE_COLLECTION,
                "--retrieval-top-k", str(_adaptive_top_k()),
                "--retrieval-out-jsonl", str(retrieval_out),
            ],
            extra_env=_embedding_env(),
        )
        steps.append({"step": "Retrieve Context", "ok": code == 0, "log": out[-2000:]})
        if code != 0:
            return jsonify(success=False, error="Context retrieval failed.", steps=steps)
    else:
        # No lectures — create empty retrieval file so grading can proceed
        retrieval_out.parent.mkdir(parents=True, exist_ok=True)
        retrieval_out.write_text("")
        steps.append({"step": "Lectures", "ok": True, "log": "No lecture chunks found; grading without RAG context."})

    # ── 5. Grade ──
    grade_args = _cli() + [
        "--mode", "grade",
        "--chunks-jsonl", str(describe_dir / "chunks.jsonl"),
        "--retrieval-out-jsonl", str(retrieval_out),
        "--output-root", str(OUTPUT_ROOT),
        "--run-id", run_id,
        "--grading-provider", grade_provider,
        "--grading-model", grade_model,
        "--student-path", student_path.name,
    ]
    if assignment_path and assignment_path.exists():
        grade_args += ["--assignment-file", str(assignment_path)]
    if rubric_path and rubric_path.exists():
        grade_args += ["--rubric-file", str(rubric_path)]

    code, out = _run(grade_args)
    steps.append({"step": "Grade Submission", "ok": code == 0, "log": out[-2000:]})
    if code != 0:
        return jsonify(success=False, error="Grading failed.", steps=steps, log=out[-1000:])

    # ── Load result ──
    grades_path = run_root / "grading" / "grades.json"
    if not grades_path.exists():
        return jsonify(success=False, error="grades.json not produced.", steps=steps)

    grades = json.loads(grades_path.read_text(encoding="utf-8"))
    # Attach grading context so UI can show what was used
    # Use original filenames (strip UUID prefix added by _safe_upload_name)
    def _original_name(path: Path | None, fallback: str = "—") -> str:
        if not path:
            return fallback
        # Uploaded files: {uuid8}_{original} — strip the prefix
        name = path.name
        if len(name) > 9 and name[8] == "_" and all(c in "0123456789abcdef" for c in name[:8]):
            return name[9:]
        return name

    grades["_grading_context"] = {
        "student_file": original_student_name,
        "rubric_file": _original_name(rubric_path),
        "assignment_file": _original_name(assignment_path),
        "describe_provider": describe_provider,
        "describe_model": describe_model,
        "grade_provider": grade_provider,
        "grade_model": grade_model,
    }
    # Auto-generate PDF report (FIFO queue, max 10) — failure is non-fatal
    try:
        pdf_path = _generate_grade_pdf(grades, run_id)
        pdf_filename = pdf_path.name if pdf_path else None
    except Exception:
        pdf_filename = None

    return jsonify(success=True, grades=grades, run_id=run_id, steps=steps,
                   pdf_report=pdf_filename)


@app.route("/api/describe", methods=["POST"])
def api_describe():
    student = (
        request.files.get("student_pdf")
        or request.files.get("student_submission")
        or request.files.get("student_file")
    )
    if not student or not student.filename:
        return jsonify(success=False, error="Please upload a student submission file."), 400
    if not _is_allowed_ext(student.filename, STUDENT_ALLOWED_EXTS):
        return jsonify(success=False, error="Invalid student file type. Allowed: PDF, PPTX, XLSX."), 400

    provider = request.form.get("describe_provider", "openai")
    model = request.form.get("describe_model") or PROVIDERS.get(provider, {}).get("model", "gpt-4o-mini")

    key_env_name = PROVIDER_API_KEY_ENV.get(provider)
    if key_env_name and not os.getenv(key_env_name):
        return jsonify(
            success=False,
            error=f"{key_env_name} is not set. Add it to the project .env file or export it before starting Flask.",
        ), 400

    run_id = f"web_describe_{datetime.now().strftime('%Y%m%d_%H%M%S')}_{uuid.uuid4().hex[:4]}"
    run_root = OUTPUT_ROOT / run_id
    upload_dir = run_root / "uploads"
    upload_dir.mkdir(parents=True, exist_ok=True)

    student_path = upload_dir / _safe_upload_name(student.filename)
    student.save(str(student_path))

    safe_model = _safe(model)
    extract_dir = run_root / "extract"
    describe_dir = run_root / f"describe_student_{provider}_{safe_model}"
    steps: list[dict] = []

    code, out = _run(
        _cli() + [
            "--mode", "extract",
            "--data-dir", str(upload_dir),
            "--output-root", str(OUTPUT_ROOT),
            "--run-id", run_id,
        ]
    )
    steps.append({"step": "Extract Submission", "ok": code == 0, "log": out[-2000:]})
    if code != 0:
        return jsonify(success=False, error="Submission extraction failed.", steps=steps, log=out[-1000:])

    code, out = _run(
        _cli() + [
            "--mode", "describe",
            "--extract-dir", str(extract_dir),
            "--describe-dir", str(describe_dir),
            "--vision-provider", provider,
            "--vision-model", model,
            "--prompt-version", "verbose_v2",
        ]
    )
    steps.append({"step": "Describe Submission", "ok": code == 0, "log": out[-2000:]})
    if code != 0:
        return jsonify(success=False, error="Describe failed.", steps=steps, log=out[-1000:])

    summary_path = describe_dir / "summary.json"
    if not summary_path.exists():
        return jsonify(success=False, error="summary.json not produced.", steps=steps), 500

    summary = json.loads(summary_path.read_text(encoding="utf-8"))
    return jsonify(
        success=True,
        run_id=run_id,
        steps=steps,
        describe={
            "summary": summary,
            "summary_path": str(summary_path),
            "chunks_jsonl": str(describe_dir / "chunks.jsonl"),
            "describe_dir": str(describe_dir),
        },
    )


@app.route("/api/grade-existing", methods=["POST"])
def api_grade_existing():
    provider = request.form.get("provider", "openai")
    model = request.form.get("model") or PROVIDERS.get(provider, {}).get("model", "gpt-4o-mini")
    student_filter = (request.form.get("existing_student_path") or "").strip()
    chunks_jsonl_raw = (request.form.get("existing_chunks_jsonl") or "").strip()
    retrieval_jsonl_raw = (request.form.get("existing_retrieval_jsonl") or "").strip()

    if not chunks_jsonl_raw:
        return jsonify(success=False, error="Existing chunks.jsonl path is required."), 400
    if not retrieval_jsonl_raw:
        return jsonify(success=False, error="Existing retrieval.jsonl path is required."), 400
    if not student_filter:
        return jsonify(success=False, error="Student path filter is required for grade-only runs."), 400

    chunks_jsonl = _resolve_project_file(chunks_jsonl_raw, allowed_exts={".jsonl"})
    retrieval_jsonl = _resolve_project_file(retrieval_jsonl_raw, allowed_exts={".jsonl"})
    if chunks_jsonl is None:
        return jsonify(success=False, error="Existing chunks.jsonl path is invalid or outside the project."), 400
    if retrieval_jsonl is None:
        return jsonify(success=False, error="Existing retrieval.jsonl path is invalid or outside the project."), 400

    key_env_name = PROVIDER_API_KEY_ENV.get(provider)
    if key_env_name and not os.getenv(key_env_name):
        return jsonify(
            success=False,
            error=f"{key_env_name} is not set. Add it to the project .env file or export it before starting Flask.",
        ), 400

    run_id = f"web_gradeexisting_{datetime.now().strftime('%Y%m%d_%H%M%S')}_{uuid.uuid4().hex[:4]}"
    run_root = OUTPUT_ROOT / run_id
    run_root.mkdir(parents=True, exist_ok=True)

    rubric_path, assignment_path, support_error = _collect_supporting_files(run_root)
    if support_error:
        return jsonify(success=False, error=support_error), 400

    grade_args = _cli() + [
        "--mode", "grade",
        "--chunks-jsonl", str(chunks_jsonl),
        "--retrieval-out-jsonl", str(retrieval_jsonl),
        "--output-root", str(OUTPUT_ROOT),
        "--run-id", run_id,
        "--grading-provider", provider,
        "--grading-model", model,
        "--student-path", student_filter,
    ]
    if assignment_path and assignment_path.exists():
        grade_args += ["--assignment-file", str(assignment_path)]
    if rubric_path and rubric_path.exists():
        grade_args += ["--rubric-file", str(rubric_path)]

    code, out = _run(grade_args)
    steps = [{"step": "Grade Existing Describe Output", "ok": code == 0, "log": out[-2000:]}]
    if code != 0:
        return jsonify(success=False, error="Grade-only run failed.", steps=steps, log=out[-1000:]), 500

    grades_path = run_root / "grading" / "grades.json"
    if not grades_path.exists():
        return jsonify(success=False, error="grades.json not produced.", steps=steps), 500

    grades = json.loads(grades_path.read_text(encoding="utf-8"))
    grades["_grading_context"] = {
        "student_file": student_filter,
        "rubric_file": rubric_path.name if rubric_path else "—",
        "assignment_file": assignment_path.name if assignment_path else "—",
        "describe_provider": "—",
        "describe_model": "—",
        "grade_provider": provider,
        "grade_model": model,
    }
    return jsonify(success=True, grades=grades, run_id=run_id, steps=steps)


# ── History endpoint — list past runs ──────────────────────────────────────
@app.route("/api/history")
def api_history():
    runs: list[dict] = []
    if OUTPUT_ROOT.exists():
        for d in sorted(OUTPUT_ROOT.iterdir(), reverse=True):
            if d.is_dir() and d.name.startswith("web_"):
                gp = d / "grading" / "grades.json"
                if gp.exists():
                    try:
                        g = json.loads(gp.read_text(encoding="utf-8"))
                        # run_id format: web_YYYYMMDD_HHMMSS_xxxx
                        # Extract readable timestamp from dir name
                        parts = d.name.split("_")  # ['web', 'YYYYMMDD', 'HHMMSS', 'xxxx']
                        if len(parts) >= 3:
                            ts = f"{parts[1]} {parts[2][:2]}:{parts[2][2:4]}:{parts[2][4:]}"
                        else:
                            ts = d.name
                        runs.append({
                            "run_id": d.name,
                            "student_file": g.get("student_file", "unknown"),
                            "score": g.get("overall_score", 0),
                            "model": g.get("grading_model", ""),
                            "timestamp": ts,
                        })
                    except Exception:
                        pass
            if len(runs) >= 20:
                break
    return jsonify(runs)


# ── Export grades as CSV ────────────────────────────────────────────────────
@app.route("/api/export-csv")
def api_export_csv():
    """Export all graded runs as a CSV file for LMS submission."""
    rows: list[dict] = []
    if OUTPUT_ROOT.exists():
        for d in sorted(OUTPUT_ROOT.iterdir()):
            if not (d.is_dir() and d.name.startswith("web_")):
                continue
            gp = d / "grading" / "grades.json"
            if not gp.exists():
                continue
            try:
                g = json.loads(gp.read_text(encoding="utf-8"))
                criterion_scores = g.get("criterion_scores", [])
                row: dict = {
                    "run_id": d.name,
                    "student_file": g.get("student_file", ""),
                    "overall_score": g.get("overall_score", ""),
                    "grading_model": g.get("grading_model", ""),
                    "confidence": g.get("confidence", ""),
                    "overall_feedback": g.get("overall_feedback", ""),
                }
                for cs in criterion_scores:
                    cid = cs.get("criterion_id", cs.get("criterion_name", "?"))
                    row[f"{cid}_awarded"] = cs.get("awarded_points", "")
                    row[f"{cid}_max"] = cs.get("max_points", "")
                rows.append(row)
            except Exception:
                pass

    if not rows:
        return jsonify(success=False, error="No graded runs found to export."), 404

    # Build CSV in memory
    fieldnames = list(rows[0].keys())
    # Make sure all rows have same keys
    for r in rows[1:]:
        for k in r:
            if k not in fieldnames:
                fieldnames.append(k)

    output = io.StringIO()
    writer = csv.DictWriter(output, fieldnames=fieldnames, extrasaction="ignore")
    writer.writeheader()
    writer.writerows(rows)

    filename = f"grades_{datetime.now().strftime('%Y%m%d_%H%M%S')}.csv"
    return Response(
        output.getvalue(),
        mimetype="text/csv",
        headers={"Content-Disposition": f"attachment; filename={filename}"},
    )


# ── Web Link Indexer ─────────────────────────────────────────────────────────
@app.route("/api/add-web-links", methods=["POST"])
def api_add_web_links():
    """
    Fetch URLs, summarize each page to 2-3 paragraphs using an LLM,
    append the summaries to the lecture chunks file, and rebuild the Chroma index.
    """
    import hashlib, textwrap
    try:
        import requests as _requests
    except ImportError:
        return jsonify(success=False, error="'requests' library not installed. Run: pip install requests"), 500
    try:
        from bs4 import BeautifulSoup as _BS
    except ImportError:
        return jsonify(success=False, error="'beautifulsoup4' not installed."), 500

    data = request.get_json(silent=True) or {}
    urls: list[str] = [u.strip() for u in (data.get("urls") or []) if u.strip()]
    provider: str = str(data.get("provider") or "openai").lower()
    model: str = str(data.get("model") or "").strip()

    if not urls:
        return jsonify(success=False, error="No URLs provided."), 400

    api_key = None
    if provider == "openai":
        api_key = os.getenv("OPENAI_API_KEY")
        if not model:
            model = "gpt-4o-mini"   # cheap but good for summarisation
    elif provider == "gemini":
        api_key = os.getenv("GEMINI_API_KEY") or os.getenv("GOOGLE_API_KEY")
        if not model:
            model = "gemini-2.0-flash"
    elif provider == "anthropic":
        api_key = os.getenv("ANTHROPIC_API_KEY")
        if not model:
            model = "claude-haiku-4-5"
    else:
        return jsonify(success=False, error=f"Unknown provider: {provider}"), 400

    if not api_key:
        return jsonify(success=False, error=f"No API key found for provider '{provider}'."), 400

    SUMMARY_PROMPT = (
        "You are a teaching assistant summarizing a web page for students in a graduate course. "
        "Read the page content below and write a concise summary of 2-3 paragraphs. "
        "Focus only on the educational concepts, facts, definitions, and processes relevant to the topic. "
        "Ignore navigation menus, advertisements, author bios, and unrelated content. "
        "Write in plain prose — no bullet points, no headings."
    )

    def _fetch_page_text(url: str) -> tuple[str, str]:
        """Fetch a URL and return (title, extracted_text). Raises on failure."""
        headers = {"User-Agent": "Mozilla/5.0 (compatible; GradeAI-Bot/1.0)"}
        resp = _requests.get(url, headers=headers, timeout=20, allow_redirects=True)
        resp.raise_for_status()
        soup = _BS(resp.text, "html.parser")
        # Extract title
        title_tag = soup.find("title")
        title = title_tag.get_text(" ", strip=True) if title_tag else url
        # Remove chrome
        for tag in soup.find_all(["script", "style", "nav", "header", "footer",
                                   "aside", "noscript", "button", "form"]):
            tag.decompose()
        # Get main text
        tags = soup.find_all(["h1","h2","h3","h4","p","li","blockquote","td","th","pre","code","dt","dd"])
        lines = [t.get_text(" ", strip=True) for t in tags]
        text = "\n".join(l for l in lines if len(l) > 20)
        return title, text[:12000]   # cap input to ~12k chars to stay within token limits

    def _summarize(title: str, text: str) -> str:
        """Call the chosen LLM to summarize the page text."""
        user_msg = f"Page title: {title}\n\nPage content:\n{text}"
        if provider == "openai":
            import urllib.request, urllib.error
            body = json.dumps({
                "model": model,
                "messages": [
                    {"role": "system", "content": SUMMARY_PROMPT},
                    {"role": "user",   "content": user_msg},
                ],
                "max_tokens": 600,
                "temperature": 0.3,
            }).encode()
            req = urllib.request.Request(
                "https://api.openai.com/v1/chat/completions",
                data=body,
                headers={"Authorization": f"Bearer {api_key}", "Content-Type": "application/json"},
            )
            with urllib.request.urlopen(req, timeout=60) as r:
                resp_data = json.loads(r.read())
            return resp_data["choices"][0]["message"]["content"].strip()

        elif provider == "gemini":
            import urllib.request
            url_api = f"https://generativelanguage.googleapis.com/v1beta/models/{model}:generateContent?key={api_key}"
            body = json.dumps({
                "contents": [{"parts": [{"text": SUMMARY_PROMPT + "\n\n" + user_msg}]}],
                "generationConfig": {"maxOutputTokens": 600, "temperature": 0.3},
            }).encode()
            req = urllib.request.Request(url_api, data=body, headers={"Content-Type": "application/json"})
            with urllib.request.urlopen(req, timeout=60) as r:
                resp_data = json.loads(r.read())
            return resp_data["candidates"][0]["content"]["parts"][0]["text"].strip()

        elif provider == "anthropic":
            import urllib.request
            body = json.dumps({
                "model": model,
                "system": SUMMARY_PROMPT,
                "messages": [{"role": "user", "content": user_msg}],
                "max_tokens": 600,
            }).encode()
            req = urllib.request.Request(
                "https://api.anthropic.com/v1/messages",
                data=body,
                headers={
                    "x-api-key": api_key,
                    "anthropic-version": "2023-06-01",
                    "Content-Type": "application/json",
                },
            )
            with urllib.request.urlopen(req, timeout=60) as r:
                resp_data = json.loads(r.read())
            return resp_data["content"][0]["text"].strip()

        return ""

    from core.chunking import chunk_text, sha1_id, make_sort_key

    results = []
    new_chunks: list[dict] = []

    for url in urls:
        try:
            title, page_text = _fetch_page_text(url)
            if len(page_text) < 50:
                results.append({"url": url, "ok": False, "error": "Page too short or empty after extraction."})
                continue

            summary = _summarize(title, page_text)
            if not summary:
                results.append({"url": url, "ok": False, "error": "LLM returned empty summary."})
                continue

            # Build chunk(s) from the summary
            slug = re.sub(r"[^a-z0-9]+", "_", url.lower())[:60]
            url_chunk_count = 0
            for ci, piece in enumerate(chunk_text(summary, 1800, 140), 1):
                if len(piece) < 20:
                    continue
                cid = sha1_id(f"web_link|{url}|chunk={ci}")
                new_chunks.append({
                    "id": cid,
                    "content": piece,
                    "metadata": {
                        "filename": slug + ".web",
                        "source_path": url,
                        "source_type": "lecture",
                        "format": "web_summary",
                        "page_number": 1,
                        "block_index": ci,
                        "sort_key": make_sort_key(1, ci),
                        "document_order": ci,
                        "content_type": "text",
                        "element_tag": "web_summary",
                        "web_url": url,
                        "web_title": title,
                        "chunk_index_in_block": ci,
                        "image_quality": None, "image_width_px": None,
                        "image_height_px": None, "image_total_pixels": None,
                        "image_aspect_ratio": None, "is_tiled": None,
                        "tile_count": None, "gpt4o_called": False, "quality_warning": None,
                    },
                })
                url_chunk_count += 1

            results.append({"url": url, "ok": True, "title": title, "summary_chars": len(summary), "chunks": url_chunk_count})

        except Exception as exc:
            results.append({"url": url, "ok": False, "error": str(exc)[:200]})

    if not new_chunks:
        return jsonify(success=False, error="No chunks generated from provided URLs.", results=results)

    # Append new chunks to the lecture chunks file
    chunks_path = DEFAULT_LECTURE_CHUNKS
    if not chunks_path.exists():
        chunks_path.parent.mkdir(parents=True, exist_ok=True)
        chunks_path.write_text("", encoding="utf-8")

    # De-duplicate: don't re-add chunks already in the file
    existing_ids: set[str] = set()
    for line in chunks_path.read_text(encoding="utf-8").splitlines():
        if line.strip():
            try:
                existing_ids.add(json.loads(line)["id"])
            except Exception:
                pass

    added_count = 0
    with chunks_path.open("a", encoding="utf-8") as f:
        for chunk in new_chunks:
            if chunk["id"] not in existing_ids:
                f.write(json.dumps(chunk, ensure_ascii=True) + "\n")
                added_count += 1

    if added_count == 0:
        return jsonify(success=True, message="All URLs already indexed (no new chunks added).", results=results)

    # Rebuild Chroma index with the new chunks
    chroma_dir = _shared_chroma_dir()
    chroma_dir.mkdir(parents=True, exist_ok=True)
    code, out = _run(
        _cli() + [
            "--mode", "index",
            "--chunks-jsonl", str(chunks_path),
            "--chroma-path", str(chroma_dir),
            "--chroma-collection", SHARED_LECTURE_COLLECTION,
            "--output-root", str(OUTPUT_ROOT),
            "--run-id", "shared_lecture_index",
        ],
        extra_env=_embedding_env(),
    )

    if code != 0:
        return jsonify(
            success=False,
            error="Chunks saved but Chroma re-index failed.",
            added_chunks=added_count,
            results=results,
            log=out[-1000:],
        ), 500

    return jsonify(
        success=True,
        message=f"Indexed {added_count} new chunk(s) from {sum(1 for r in results if r['ok'])} URL(s).",
        added_chunks=added_count,
        results=results,
    )



# ── Web Link Summarizer (demo: fetch + summarise, no indexing) ───────────────
@app.route("/api/summarize-web-link", methods=["POST"])
def api_summarize_web_link():
    """
    Fetch a URL, extract its text, and summarize it to 2-3 paragraphs using an LLM.
    Returns the title, raw extracted text preview, and the full AI summary.
    No indexing — purely for demonstration of extraction quality.
    """
    try:
        import requests as _requests
    except ImportError:
        return jsonify(success=False, error="'requests' library not installed."), 500
    try:
        from bs4 import BeautifulSoup as _BS
    except ImportError:
        return jsonify(success=False, error="'beautifulsoup4' not installed."), 500

    data = request.get_json(silent=True) or {}
    url: str = str(data.get("url") or "").strip()
    provider: str = str(data.get("provider") or "openai").lower()
    model: str = str(data.get("model") or "").strip()

    if not url:
        return jsonify(success=False, error="No URL provided."), 400

    api_key = None
    if provider == "openai":
        api_key = os.getenv("OPENAI_API_KEY")
        if not model:
            model = "gpt-4o-mini"
    elif provider == "gemini":
        api_key = os.getenv("GEMINI_API_KEY") or os.getenv("GOOGLE_API_KEY")
        if not model:
            model = "gemini-2.0-flash"
    elif provider == "anthropic":
        api_key = os.getenv("ANTHROPIC_API_KEY")
        if not model:
            model = "claude-haiku-4-5"
    else:
        return jsonify(success=False, error=f"Unknown provider: {provider}"), 400

    if not api_key:
        return jsonify(success=False, error=f"No API key found for provider '{provider}'."), 400

    SUMMARY_PROMPT = (
        "You are a teaching assistant summarizing a web page for students in a graduate course. "
        "Read the page content below and write a concise summary of 2-3 paragraphs. "
        "Focus only on the educational concepts, facts, definitions, and processes relevant to the topic. "
        "Ignore navigation menus, advertisements, author bios, and unrelated content. "
        "Write in plain prose — no bullet points, no headings."
    )

    def _fetch_page_text(u: str) -> tuple[str, str, list[str]]:
        """Returns (title, full_text_for_llm, extracted_lines_preview)."""
        headers = {"User-Agent": "Mozilla/5.0 (compatible; GradeAI-Bot/1.0)"}
        resp = _requests.get(u, headers=headers, timeout=20, allow_redirects=True)
        resp.raise_for_status()
        soup = _BS(resp.text, "html.parser")
        title_tag = soup.find("title")
        title = title_tag.get_text(" ", strip=True) if title_tag else u
        for tag in soup.find_all(["script", "style", "nav", "header", "footer",
                                   "aside", "noscript", "button", "form"]):
            tag.decompose()
        tags = soup.find_all(["h1","h2","h3","h4","p","li","blockquote","td","th","pre","code","dt","dd"])
        lines = [t.get_text(" ", strip=True) for t in tags if len(t.get_text(" ", strip=True)) > 20]
        text = "\n".join(lines)
        return title, text[:12000], lines[:30]   # preview = first 30 meaningful lines

    def _summarize(title: str, text: str) -> str:
        user_msg = f"Page title: {title}\n\nPage content:\n{text}"
        if provider == "openai":
            import urllib.request
            body = json.dumps({
                "model": model,
                "messages": [
                    {"role": "system", "content": SUMMARY_PROMPT},
                    {"role": "user",   "content": user_msg},
                ],
                "max_tokens": 700, "temperature": 0.3,
            }).encode()
            req = urllib.request.Request(
                "https://api.openai.com/v1/chat/completions", data=body,
                headers={"Authorization": f"Bearer {api_key}", "Content-Type": "application/json"},
            )
            with urllib.request.urlopen(req, timeout=60) as r:
                return json.loads(r.read())["choices"][0]["message"]["content"].strip()

        elif provider == "gemini":
            import urllib.request
            url_api = f"https://generativelanguage.googleapis.com/v1beta/models/{model}:generateContent?key={api_key}"
            body = json.dumps({
                "contents": [{"parts": [{"text": SUMMARY_PROMPT + "\n\n" + user_msg}]}],
                "generationConfig": {"maxOutputTokens": 700, "temperature": 0.3},
            }).encode()
            req = urllib.request.Request(url_api, data=body, headers={"Content-Type": "application/json"})
            with urllib.request.urlopen(req, timeout=60) as r:
                return json.loads(r.read())["candidates"][0]["content"]["parts"][0]["text"].strip()

        elif provider == "anthropic":
            import urllib.request
            body = json.dumps({
                "model": model, "system": SUMMARY_PROMPT,
                "messages": [{"role": "user", "content": user_msg}],
                "max_tokens": 700,
            }).encode()
            req = urllib.request.Request(
                "https://api.anthropic.com/v1/messages", data=body,
                headers={"x-api-key": api_key, "anthropic-version": "2023-06-01",
                         "Content-Type": "application/json"},
            )
            with urllib.request.urlopen(req, timeout=60) as r:
                return json.loads(r.read())["content"][0]["text"].strip()
        return ""

    try:
        title, page_text, preview_lines = _fetch_page_text(url)
        if len(page_text) < 50:
            return jsonify(success=False, error="Page too short or empty after extraction.")
        summary = _summarize(title, page_text)
        if not summary:
            return jsonify(success=False, error="LLM returned empty summary.")
        return jsonify(
            success=True,
            url=url,
            title=title,
            extracted_chars=len(page_text),
            extracted_lines=len(preview_lines),
            raw_preview=preview_lines,   # first ~30 lines extracted from the page
            summary=summary,
            model_used=model,
            provider=provider,
        )
    except Exception as exc:
        return jsonify(success=False, error=str(exc)[:300])


# ── Lecture Index Inspector ──────────────────────────────────────────────────
@app.route("/api/lecture-index-stats")
def api_lecture_index_stats():
    """Return detailed stats about what's in the lecture index for the professor."""
    from collections import defaultdict

    if not DEFAULT_LECTURE_CHUNKS.exists():
        return jsonify(success=False, error="No lecture chunks file found.")

    lines = [l for l in DEFAULT_LECTURE_CHUNKS.read_text(encoding="utf-8").splitlines() if l.strip()]
    total_chunks = len(lines)

    # Aggregate per file
    files: dict = defaultdict(lambda: {
        "filename": "", "format": "", "chunks": 0,
        "tag_counts": defaultdict(int), "sample": "",
        "is_web": False, "web_url": "", "web_title": "",
    })
    format_counts: dict = defaultdict(int)
    tag_counts: dict = defaultdict(int)
    total_chars = 0

    for line in lines:
        try:
            c = json.loads(line)
        except Exception:
            continue
        meta = c.get("metadata", {}) or {}
        fname = meta.get("filename", "unknown")
        fmt = meta.get("format", "unknown")
        tag = meta.get("element_tag") or fmt
        content = str(c.get("content", ""))

        f = files[fname]
        f["filename"] = fname
        f["format"] = fmt
        f["chunks"] += 1
        f["tag_counts"][tag] = f["tag_counts"].get(tag, 0) + 1
        if not f["sample"]:
            f["sample"] = content[:200]
        if fmt == "web_summary":
            f["is_web"] = True
            f["web_url"] = meta.get("web_url", "")
            f["web_title"] = meta.get("web_title", "")

        format_counts[fmt] += 1
        tag_counts[tag] = tag_counts.get(tag, 0) + 1
        total_chars += len(content)

    # Sort files: web summaries last, then by chunk count desc
    sorted_files = sorted(
        files.values(),
        key=lambda x: (x["is_web"], -x["chunks"])
    )
    # Make tag_counts serialisable (defaultdict → dict)
    for f in sorted_files:
        f["tag_counts"] = dict(f["tag_counts"])

    return jsonify(
        success=True,
        total_chunks=total_chunks,
        total_files=len(files),
        total_chars=total_chars,
        avg_chunk_chars=round(total_chars / max(1, total_chunks)),
        format_breakdown=dict(format_counts),
        tag_breakdown=dict(tag_counts),
        chroma_ready=_shared_chroma_ready(),
        embedding_provider=_embedding_provider(),
        files=sorted_files,
    )


@app.route("/api/lecture-search")
def api_lecture_search():
    """Test RAG retrieval: given a query, return the top lecture chunks retrieved."""
    query = request.args.get("q", "").strip()
    top_k = int(request.args.get("k", "5"))
    if not query:
        return jsonify(success=False, error="No query provided.")
    if not _shared_chroma_ready():
        return jsonify(success=False, error="Lecture index not ready. Click 'Build Lecture Index' first.")

    try:
        import chromadb
        sys.path.insert(0, str(SCRIPTS_DIR))
        from storage.chroma_store import _build_embedding_function

        chroma_path = str(_shared_chroma_dir())
        client = chromadb.PersistentClient(path=chroma_path)
        embed_fn = _build_embedding_function()
        if embed_fn:
            col = client.get_or_create_collection("lecture_v1", embedding_function=embed_fn)
        else:
            col = client.get_or_create_collection("lecture_v1")

        result = col.query(
            query_texts=[query],
            n_results=min(top_k, col.count()),
            include=["documents", "metadatas", "distances"],
        )
        hits = []
        docs = (result.get("documents") or [[]])[0]
        metas = (result.get("metadatas") or [[]])[0]
        dists = (result.get("distances") or [[]])[0]
        for doc, meta, dist in zip(docs, metas, dists):
            hits.append({
                "content": doc[:400],
                "filename": (meta or {}).get("filename", ""),
                "element_tag": (meta or {}).get("element_tag", ""),
                "distance": round(float(dist), 4),
                "relevance_pct": round(max(0, (1 - float(dist) / 2)) * 100, 1),
            })
        return jsonify(success=True, query=query, hits=hits)
    except Exception as exc:
        return jsonify(success=False, error=str(exc))


# ── Status endpoint ─────────────────────────────────────────────────────────
@app.route("/api/status")
def api_status():
    return jsonify({
        "lecture_chunks_exist": DEFAULT_LECTURE_CHUNKS.exists(),
        "shared_chroma_ready": _shared_chroma_ready(),
        "shared_chroma_path": str(_shared_chroma_dir()),
        "embedding_provider": _embedding_provider(),
        "lecture_chunks_path": str(DEFAULT_LECTURE_CHUNKS),
    })


@app.route("/api/library/save-assignment", methods=["POST"])
def api_library_save_assignment():
    """Save an uploaded assignment description file to the library.
    Filename MUST match assignment_N.ext or Assignment_N.ext (N = integer).
    """
    f = request.files.get("file")
    if not f or not f.filename:
        return jsonify(success=False, error="No file uploaded."), 400
    ext = Path(f.filename).suffix.lower()
    if ext not in SUPPORT_ALLOWED_EXTS:
        return jsonify(success=False, error=f"Invalid file type. Allowed: {', '.join(sorted(SUPPORT_ALLOWED_EXTS))}"), 400
    # Validate naming convention
    if not re.match(r"(?i)^assignment[_\-]?\d+\.", f.filename):
        return jsonify(success=False,
            error="Filename must follow the format: assignment_1.pdf, assignment_2.docx, etc."), 400
    # Canonical saved name: lower-case, underscore
    num = _assignment_number(f.filename)
    canonical = f"assignment_{num}{ext}"
    LIBRARY_ASSIGNMENTS_DIR.mkdir(parents=True, exist_ok=True)
    dest = LIBRARY_ASSIGNMENTS_DIR / canonical
    if dest.exists():
        return jsonify(success=False,
            error=f"Assignment - {num} - Description already exists in the library. Delete it first to replace."), 409
    f.save(str(dest))
    return jsonify(success=True,
                   display_name=_display_name(dest.name, kind="description"),
                   assignment_number=num,
                   filename=dest.name,
                   path=str(dest))


@app.route("/api/library/save-quiz", methods=["POST"])
def api_library_save_quiz():
    """Save an uploaded quiz description file to the library.
    Filename MUST match quiz_N.ext or Quiz_N.ext (N = integer).
    """
    f = request.files.get("file")
    if not f or not f.filename:
        return jsonify(success=False, error="No file uploaded."), 400
    ext = Path(f.filename).suffix.lower()
    if ext not in SUPPORT_ALLOWED_EXTS:
        return jsonify(success=False, error=f"Invalid file type. Allowed: {', '.join(sorted(SUPPORT_ALLOWED_EXTS))}"), 400
    if not re.match(r"(?i)^quiz[_\-]?\d+\.", f.filename):
        return jsonify(success=False,
            error="Filename must follow the format: quiz_1.pdf, quiz_2.docx, etc."), 400
    num = _quiz_number(f.filename)
    canonical = f"quiz_{num}{ext}"
    LIBRARY_QUIZZES_DIR.mkdir(parents=True, exist_ok=True)
    dest = LIBRARY_QUIZZES_DIR / canonical
    if dest.exists():
        return jsonify(success=False,
            error=f"Quiz - {num} - Description already exists in the library. Delete it first to replace."), 409
    f.save(str(dest))
    return jsonify(success=True,
                   display_name=_display_name(dest.name, kind="description"),
                   quiz_number=num,
                   filename=dest.name,
                   path=str(dest))


@app.route("/api/library/save-rubric", methods=["POST"])
def api_library_save_rubric():
    """Save an uploaded rubric file to the library."""
    f = request.files.get("file")
    if not f or not f.filename:
        return jsonify(success=False, error="No file uploaded."), 400
    ext = Path(f.filename).suffix.lower()
    if ext not in RUBRIC_ALLOWED_EXTS:
        return jsonify(success=False, error=f"Invalid file type. Allowed: {', '.join(sorted(RUBRIC_ALLOWED_EXTS))}"), 400
    # Try to extract assignment/quiz number for canonical naming
    num = _assignment_number(f.filename)
    if num:
        canonical = f"assignment_{num}_rubric{ext}"
    elif (qnum := _quiz_number(f.filename)):
        canonical = f"quiz_{qnum}_rubric{ext}"
    else:
        safe = re.sub(r"[^\w\-]", "_", Path(f.filename).stem).strip("_") or "rubric"
        canonical = f"{safe}{ext}"
    LIBRARY_RUBRICS_DIR.mkdir(parents=True, exist_ok=True)
    dest = LIBRARY_RUBRICS_DIR / canonical
    if dest.exists():
        return jsonify(success=False,
            error=f"'{_display_name(canonical, kind='rubric')}' already exists. Delete it first to replace."), 409
    f.save(str(dest))
    return jsonify(success=True,
                   display_name=_display_name(dest.name, kind="rubric"),
                   assignment_number=num,
                   filename=dest.name,
                   path=str(dest))


@app.route("/api/library/delete", methods=["POST"])
def api_library_delete():
    """Delete a file from the library."""
    data = request.get_json(force=True, silent=True) or {}
    file_type = data.get("type", "")
    filename   = data.get("filename", "")
    if not filename or ".." in filename or "/" in filename or "\\" in filename:
        return jsonify(success=False, error="Invalid filename."), 400
    if file_type == "assignment":
        path = LIBRARY_ASSIGNMENTS_DIR / filename
    elif file_type == "quiz":
        path = LIBRARY_QUIZZES_DIR / filename
    elif file_type == "rubric":
        path = LIBRARY_RUBRICS_DIR / filename
    else:
        return jsonify(success=False, error="Unknown type. Use 'assignment', 'quiz', or 'rubric'."), 400
    if not path.exists():
        return jsonify(success=False, error="File not found."), 404
    path.unlink()
    return jsonify(success=True)


@app.route("/api/library/list", methods=["GET"])
def api_library_list():
    """Return current library contents for dropdown refresh."""
    rubrics = (
        _discover_files(DEFAULT_RUBRIC_DIR, kind="rubric") +
        _discover_files(LIBRARY_RUBRICS_DIR, exts=RUBRIC_ALLOWED_EXTS, kind="rubric")
    )
    assignments = (
        _discover_files(PROJECT_ROOT / "assignments") +
        _discover_files(LIBRARY_ASSIGNMENTS_DIR)
    )
    quizzes = _discover_files(LIBRARY_QUIZZES_DIR)
    return jsonify(rubrics=rubrics, assignments=assignments, quizzes=quizzes)


@app.route("/api/library/preview", methods=["GET"])
def api_library_preview():
    """Return the content of a library file for the inline rubric preview panel."""
    file_type = request.args.get("type", "")
    filename  = request.args.get("filename", "")
    if not filename or ".." in filename or "/" in filename or "\\" in filename:
        return jsonify(success=False, error="Invalid filename."), 400
    if file_type == "rubric":
        path = LIBRARY_RUBRICS_DIR / filename
    elif file_type == "assignment":
        path = LIBRARY_ASSIGNMENTS_DIR / filename
    elif file_type == "quiz":
        path = LIBRARY_QUIZZES_DIR / filename
    else:
        return jsonify(success=False, error="Unknown type. Use 'rubric', 'assignment', or 'quiz'."), 400
    if not path.exists():
        return jsonify(success=False, error="File not found in library."), 404

    ext = path.suffix.lower()
    if ext == ".json":
        try:
            data = json.loads(path.read_text(encoding="utf-8"))
            return jsonify(success=True, content_type="json", data=data)
        except Exception as e:
            return jsonify(success=False, error=f"Invalid JSON: {e}"), 400
    elif ext in {".txt", ".md"}:
        text = path.read_text(encoding="utf-8", errors="ignore")
        return jsonify(success=True, content_type="text", preview=text[:4000])
    elif ext == ".docx":
        try:
            from docx import Document as _DocxDoc
            doc = _DocxDoc(str(path))
            parts: list[str] = []
            for p in doc.paragraphs:
                t = p.text.strip()
                if t:
                    parts.append(t)
            # Many rubric DOCX files store content only in tables — always extract them
            for table in doc.tables:
                for row in table.rows:
                    cells = [c.text.strip() for c in row.cells if c.text.strip()]
                    if cells:
                        parts.append(" | ".join(cells))
            return jsonify(success=True, content_type="text", preview="\n".join(parts)[:4000])
        except Exception as e:
            return jsonify(success=True, content_type="binary",
                           preview=f"DOCX preview unavailable: {e}")
    elif ext == ".pdf":
        try:
            import fitz as _fitz
            doc = _fitz.open(str(path))
            text = "\n\n".join(page.get_text() for page in doc).strip()
            return jsonify(success=True, content_type="text", preview=text[:4000])
        except Exception:
            try:
                from pypdf import PdfReader as _PdfReader
                reader = _PdfReader(str(path))
                text = "\n\n".join(
                    (page.extract_text() or "").strip() for page in reader.pages
                ).strip()
                return jsonify(success=True, content_type="text", preview=text[:4000])
            except Exception as e2:
                return jsonify(success=True, content_type="binary",
                               preview=f"PDF preview unavailable: {e2}")
    else:
        return jsonify(success=True, content_type="binary",
                       preview=f"Inline preview is not available for {ext} files.")


@app.route("/api/generate-rubric", methods=["POST"])
def api_generate_rubric():
    """Generate or enhance a rubric from assignment instructions."""
    import sys as _sys
    if str(SCRIPTS_DIR) not in _sys.path:
        _sys.path.insert(0, str(SCRIPTS_DIR))

    api_key = os.getenv("ANTHROPIC_API_KEY")
    if not api_key:
        return jsonify(success=False, error="ANTHROPIC_API_KEY is not set. Add it to the project .env file."), 400

    # ── assignment text ──
    assignment_text = (request.form.get("assignment_text") or "").strip()
    if not assignment_text and "assignment_file" in request.files:
        f = request.files["assignment_file"]
        if f and f.filename:
            if not _is_allowed_ext(f.filename, {".txt", ".pdf", ".md"}):
                return jsonify(success=False, error="Assignment file must be .txt, .pdf, or .md"), 400
            raw = f.read()
            if f.filename.lower().endswith(".pdf"):
                try:
                    import fitz
                    doc = fitz.open(stream=raw, filetype="pdf")
                    assignment_text = "\n".join(page.get_text() for page in doc).strip()
                except Exception as exc:
                    return jsonify(success=False, error=f"Failed to read PDF: {exc}"), 400
            else:
                assignment_text = raw.decode("utf-8", errors="replace").strip()

    # library path shortcut — when user picks a saved assignment from the dropdown
    if not assignment_text:
        lib_path_str = (request.form.get("assignment_library_path") or "").strip()
        if lib_path_str:
            lib_p = Path(lib_path_str)
            if lib_p.exists():
                ext_lp = lib_p.suffix.lower()
                if ext_lp == ".pdf":
                    try:
                        import fitz
                        doc = fitz.open(str(lib_p))
                        assignment_text = "\n".join(page.get_text() for page in doc).strip()
                    except Exception as exc:
                        return jsonify(success=False, error=f"Failed to read library PDF: {exc}"), 400
                elif ext_lp == ".docx":
                    try:
                        from docx import Document as _DocxDocument
                        _doc = _DocxDocument(str(lib_p))
                        parts = [p.text.strip() for p in _doc.paragraphs if p.text.strip()]
                        for tbl in _doc.tables:
                            for row in tbl.rows:
                                line = " | ".join(c.text.strip() for c in row.cells if c.text.strip())
                                if line:
                                    parts.append(line)
                        assignment_text = "\n".join(parts).strip()
                    except Exception as exc:
                        return jsonify(success=False, error=f"Failed to read library DOCX: {exc}"), 400
                else:
                    # .txt / .md and any other plain-text format
                    assignment_text = lib_p.read_text(encoding="utf-8", errors="replace").strip()
            else:
                return jsonify(success=False, error=f"Library file not found: {lib_path_str}"), 400

    if not assignment_text:
        return jsonify(success=False, error="Provide assignment_text, upload an assignment_file, or select a saved assignment."), 400

    # ── existing rubric (optional — triggers enhance mode) ──
    existing_rubric = (request.form.get("existing_rubric") or "").strip()
    if not existing_rubric and "existing_rubric_file" in request.files:
        rf = request.files["existing_rubric_file"]
        if rf and rf.filename:
            if not _is_allowed_ext(rf.filename, {".txt", ".md", ".pdf", ".docx", ".json"}):
                return jsonify(success=False, error="Existing rubric file must be .txt, .md, .pdf, .docx, or .json"), 400
            raw = rf.read()
            ext = Path(rf.filename).suffix.lower()
            if ext == ".pdf":
                try:
                    import fitz
                    doc = fitz.open(stream=raw, filetype="pdf")
                    existing_rubric = "\n".join(page.get_text() for page in doc).strip()
                except Exception as exc:
                    return jsonify(success=False, error=f"Failed to read rubric PDF: {exc}"), 400
            elif ext == ".docx":
                try:
                    import io
                    from docx import Document
                    doc = Document(io.BytesIO(raw))
                    existing_rubric = "\n".join(p.text for p in doc.paragraphs if p.text.strip())
                except Exception as exc:
                    return jsonify(success=False, error=f"Failed to read rubric DOCX: {exc}"), 400
            else:
                existing_rubric = raw.decode("utf-8", errors="replace").strip()

    instructions = (request.form.get("instructions") or "").strip()
    model = (request.form.get("model") or "claude-sonnet-4-6").strip()

    try:
        from rubric_gen.generate_rubric import generate_rubric, enhance_rubric, rubric_to_dict
    except ImportError as exc:
        return jsonify(success=False, error=f"Could not import rubric_gen: {exc}"), 500

    try:
        if existing_rubric:
            rubric = enhance_rubric(assignment_text, existing_rubric,
                                    instructions=instructions, model=model, api_key=api_key)
            mode = "enhanced"
        else:
            rubric = generate_rubric(assignment_text, instructions=instructions,
                                     model=model, api_key=api_key)
            mode = "generated"
        rubric_dict = rubric_to_dict(rubric)
        saved_path = None
        saved_display = None
        save_name = (request.form.get("save_name") or "").strip()
        if save_name:
            LIBRARY_RUBRICS_DIR.mkdir(parents=True, exist_ok=True)
            # Use canonical assignment_N_rubric.json or quiz_N_rubric.json format if number detected
            num = _assignment_number(save_name)
            qnum = _quiz_number(save_name)
            if num:
                canonical = f"assignment_{num}_rubric.json"
            elif qnum:
                canonical = f"quiz_{qnum}_rubric.json"
            elif (m := re.search(r"\d+", save_name)):
                num = m.group()
                canonical = f"assignment_{num}_rubric.json"
            else:
                safe = re.sub(r"[^\w\-]", "_", save_name).strip("_") or "generated_rubric"
                canonical = f"{safe}.json"
            dest = LIBRARY_RUBRICS_DIR / canonical
            dest.write_text(json.dumps(rubric_dict, indent=2), encoding="utf-8")
            saved_path = str(dest)
            saved_display = _display_name(canonical, kind="rubric")
        return jsonify(success=True, mode=mode, rubric=rubric_dict,
                       saved_path=saved_path, saved_display=saved_display)
    except Exception as exc:
        return jsonify(success=False, error=str(exc)), 500


# ── PDF Report Generation ────────────────────────────────────────────────────
REPORTS_DIR   = PROJECT_ROOT / "data" / "reports"
REPORTS_INDEX = REPORTS_DIR / "index.json"
MAX_REPORTS   = 10   # FIFO queue depth


def _safe_ascii(text: str) -> str:
    """
    Replace common Unicode chars with Latin-1 / ASCII equivalents so that
    fpdf2 (Helvetica, latin-1 encoding) never raises FPDFUnicodeEncodingException.
    """
    _UNICODE_MAP = {
        "\u2014": "--",   # em dash  —
        "\u2013": "-",    # en dash  –
        "\u2022": "*",    # bullet   •
        "\u2019": "'",    # right single quote  '
        "\u2018": "'",    # left single quote   '
        "\u201c": '"',    # left double quote   "
        "\u201d": '"',    # right double quote  "
        "\u2026": "...",  # ellipsis …
        "\u00b7": "-",    # middle dot · (IS latin-1 but kept safe)
        "\u00a9": "(c)",  # copyright ©
        "\u00ae": "(R)",  # registered ®
        "\u2122": "(TM)", # trademark ™
        "\u2212": "-",    # minus sign −
        "\u00d7": "x",    # multiplication x
        "\u2265": ">=",   # ≥
        "\u2264": "<=",   # ≤
        "\u2260": "!=",   # ≠
        "\u25b6": ">",    # ▶
        "\u2192": "->",   # →
        "\u2190": "<-",   # ←
    }
    for ch, repl in _UNICODE_MAP.items():
        text = text.replace(ch, repl)
    # Final safety: drop anything still outside latin-1
    return text.encode("latin-1", errors="replace").decode("latin-1")


def _generate_grade_pdf(grades: dict, run_id: str) -> Path | None:
    """
    Build a nicely-formatted PDF grade report using fpdf2.
    Returns the saved Path, or None if fpdf2 is missing or any error occurs.
    Never raises — PDF failure must not block grading.
    The FIFO queue keeps the last MAX_REPORTS PDFs; older files are deleted automatically.
    """
    try:
        from fpdf import FPDF  # fpdf2
    except ImportError:
        return None

    try:
        return _build_grade_pdf(FPDF, grades, run_id)
    except Exception as exc:
        import traceback
        print(f"[PDF] generation failed (non-fatal): {exc}\n{traceback.format_exc()}")
        return None


def _build_grade_pdf(FPDF, grades: dict, run_id: str) -> Path | None:
    """Inner implementation — called by _generate_grade_pdf inside a try/except."""
    REPORTS_DIR.mkdir(parents=True, exist_ok=True)

    student_raw  = grades.get("student_file", "unknown")
    # Strip UUID prefix added by _safe_upload_name  (8 hex chars + "_")
    student = re.sub(r"^[0-9a-f]{8}_", "", student_raw)
    score        = round(grades.get("overall_score", 0))
    total_max    = grades.get("total_max_points", 100) or 100
    feedback     = _safe_ascii(grades.get("overall_feedback", ""))
    criteria     = grades.get("criterion_details", [])
    section_cov  = grades.get("section_coverage", [])
    ctx          = grades.get("_grading_context", {})

    # Score colour thresholds — percentage-based so non-100 rubrics get correct colours
    def _score_rgb(s: float, mx: float = 100.0):
        pct = s / mx * 100 if mx else 0
        if pct >= 90: return (34, 197, 94)
        if pct >= 80: return (20, 184, 166)
        if pct >= 70: return (245, 158, 11)
        if pct >= 60: return (249, 115, 22)
        return (239, 68, 68)

    CRIT_COLORS = [
        (99, 102, 241), (16, 185, 129), (245, 158, 11),
        (239, 68, 68),  (20, 184, 166), (139, 92, 246),
        (236, 72, 153), (14, 165, 233),
    ]

    pdf = FPDF()
    pdf.set_auto_page_break(auto=True, margin=15)
    pdf.add_page()

    # Helper: shorthand for pdf.cell + pdf.ln()  (avoids deprecated ln=True)
    def _cell(w, h, txt, **kw):
        pdf.set_x(pdf.l_margin)
        pdf.cell(w, h, _safe_ascii(str(txt)), **kw)
        pdf.ln(h)

    def _row(h, txt, **kw):
        pdf.set_x(pdf.l_margin)
        pdf.cell(0, h, _safe_ascii(str(txt)), **kw)
        pdf.ln(h)

    def _mcell(h, txt, **kw):
        # Must reset X to left margin — multi_cell crashes if cursor is past page width
        pdf.set_x(pdf.l_margin)
        pdf.multi_cell(0, h, _safe_ascii(str(txt)), **kw)

    # ── Header band ──
    pdf.set_fill_color(30, 27, 75)
    pdf.rect(0, 0, 210, 32, "F")
    pdf.set_font("Helvetica", "B", 17)
    pdf.set_text_color(255, 255, 255)
    pdf.set_xy(12, 7)
    _row(10, "GradeAI Pro -- Grade Report")
    pdf.set_font("Helvetica", "", 9)
    pdf.set_xy(12, 20)
    date_str = datetime.now().strftime("%B %d, %Y  %H:%M")
    pdf.cell(0, 7, _safe_ascii(f"Student: {student}   |   Run: {run_id}   |   {date_str}"))
    pdf.set_text_color(0, 0, 0)
    pdf.ln(18)

    # ── Score hero ──
    sr, sg, sb = _score_rgb(score, total_max)
    pdf.set_fill_color(sr, sg, sb)
    pdf.set_text_color(255, 255, 255)
    pdf.set_font("Helvetica", "B", 38)
    pdf.cell(48, 22, str(score), fill=True, align="C")
    pdf.set_font("Helvetica", "", 13)
    pdf.set_text_color(100, 100, 100)
    pdf.cell(20, 22, f"/ {int(total_max)}", align="L")
    pdf.ln(26)

    # ── Context row ──
    if ctx:
        pdf.set_fill_color(248, 249, 251)
        pdf.set_draw_color(226, 232, 240)
        pdf.set_font("Helvetica", "", 8)
        pdf.set_text_color(71, 85, 105)
        ctx_line = (
            f"Assignment: {ctx.get('assignment_file','?')}   |   "
            f"Rubric: {ctx.get('rubric_file','?')}   |   "
            f"Describe: {ctx.get('describe_provider','')}/{ctx.get('describe_model','')}   |   "
            f"Grade: {ctx.get('grade_provider','')}/{ctx.get('grade_model','')}"
        )
        _row(7, ctx_line, fill=True, border=1)
        pdf.ln(4)

    # ── Overall Feedback ──
    pdf.set_text_color(0, 0, 0)
    pdf.set_font("Helvetica", "B", 11)
    _row(8, "Overall Feedback")
    pdf.set_font("Helvetica", "", 9)
    _mcell(5, feedback or "N/A")
    pdf.ln(4)

    # ── Section Coverage ──
    if section_cov:
        pdf.set_font("Helvetica", "B", 11)
        _row(8, "Section Coverage")
        pdf.set_font("Helvetica", "", 9)
        STATUS_COLORS = {"addressed": (34,197,94), "partial": (245,158,11), "missing": (239,68,68)}
        for s in section_cov:
            sid    = _safe_ascii(str(s.get("section_id", "")))
            status = str(s.get("status", "")).lower()
            cr, cg, cb = STATUS_COLORS.get(status, (148, 163, 184))
            pdf.set_fill_color(cr, cg, cb)
            pdf.set_text_color(255, 255, 255)
            pdf.cell(60, 6, f"  {sid}: {status}", fill=True, border=0)
            pdf.set_text_color(0, 0, 0)
            pdf.ln(7)
        pdf.ln(3)

    # ── Score Breakdown table ──
    if criteria:
        pdf.set_font("Helvetica", "B", 11)
        _row(8, "Score Breakdown")
        # Header row
        pdf.set_fill_color(30, 27, 75)
        pdf.set_text_color(255, 255, 255)
        pdf.set_font("Helvetica", "B", 8.5)
        pdf.cell(105, 7, "  Criterion", fill=True, border=0)
        pdf.cell(25,  7, "Awarded",    fill=True, border=0, align="C")
        pdf.cell(20,  7, "Max",        fill=True, border=0, align="C")
        pdf.cell(40,  7, "Score %",    fill=True, border=0, align="C")
        pdf.ln(7)
        pdf.set_font("Helvetica", "", 8.5)
        for i, c in enumerate(criteria):
            awarded  = c.get("awarded_points", 0)
            max_pts  = c.get("max_points", 0)
            pct      = f"{round(awarded/max_pts*100)}%" if max_pts else "?"
            cname    = _safe_ascii(str(c.get("criterion_name", ""))[:52])
            row_fill = (248, 249, 251) if i % 2 == 0 else (255, 255, 255)
            pdf.set_fill_color(*row_fill)
            pdf.set_text_color(30, 41, 59)
            pdf.cell(105, 6, f"  {cname}", fill=True, border="B")
            pdf.cell(25,  6, str(awarded), fill=True, border="B", align="C")
            pdf.cell(20,  6, str(max_pts), fill=True, border="B", align="C")
            pdf.cell(40,  6, pct,          fill=True, border="B", align="C")
            pdf.ln(6)
        pdf.ln(6)

    # ── Criterion Evidence Details ──
    if criteria:
        pdf.set_font("Helvetica", "B", 11)
        _row(8, "Criterion Evidence Details")
        pdf.ln(2)
        for i, c in enumerate(criteria):
            cr, cg, cb = CRIT_COLORS[i % len(CRIT_COLORS)]
            pdf.set_fill_color(cr, cg, cb)
            pdf.set_text_color(255, 255, 255)
            pdf.set_font("Helvetica", "B", 9.5)
            cname   = _safe_ascii(str(c.get("criterion_name", ""))[:60])
            awarded = c.get("awarded_points", 0)
            max_pts = c.get("max_points", 0)
            _row(7, f"  {cname}   [{awarded} / {max_pts} pts]", fill=True)
            pdf.set_text_color(30, 41, 59)
            pdf.set_fill_color(250, 251, 252)
            # Justification
            just = str(c.get("justification", "")).strip()
            if just:
                pdf.set_font("Helvetica", "B", 8)
                _row(5, "  Justification:", fill=True)
                pdf.set_font("Helvetica", "", 8)
                _mcell(4.5, "    " + just, fill=True)
            # Evidence refs
            evidence = [_safe_ascii(e) for e in (c.get("evidence_refs") or [])[:4] if e]
            if evidence:
                pdf.set_font("Helvetica", "B", 8)
                _row(5, "  Evidence:", fill=True)
                pdf.set_font("Helvetica", "", 8)
                for ev in evidence:
                    _mcell(4.5, f"    * {ev}", fill=True)
            # Missing items
            missing = [_safe_ascii(m) for m in (c.get("missing_items") or [])[:4] if m]
            if missing:
                pdf.set_font("Helvetica", "B", 8)
                _row(5, "  Missing / Gaps:", fill=True)
                pdf.set_font("Helvetica", "", 8)
                for m in missing:
                    _mcell(4.5, f"    * {m}", fill=True)
            pdf.ln(4)

    # ── Footer ──
    pdf.set_y(-14)
    pdf.set_font("Helvetica", "I", 7.5)
    pdf.set_text_color(148, 163, 184)
    pdf.cell(0, 6, _safe_ascii(f"Generated by GradeAI Pro  |  {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}"),
             align="C")

    # ── Save & update FIFO index ──
    safe_name = re.sub(r"[^\w\-]", "_", student)[:40]
    pdf_filename = f"{run_id}_{safe_name}.pdf"
    pdf_path = REPORTS_DIR / pdf_filename
    try:
        pdf.output(str(pdf_path))
    except Exception:
        return None

    # Load / update index
    try:
        index: list[dict] = json.loads(REPORTS_INDEX.read_text()) if REPORTS_INDEX.exists() else []
    except Exception:
        index = []

    index.insert(0, {
        "run_id":    run_id,
        "filename":  pdf_filename,
        "student":   student,
        "score":     score,
        "total_max": int(total_max),
        "timestamp": datetime.now().isoformat(),
    })
    # Evict oldest beyond MAX_REPORTS
    while len(index) > MAX_REPORTS:
        old = index.pop()
        old_path = REPORTS_DIR / old["filename"]
        if old_path.exists():
            try:
                old_path.unlink()
            except Exception:
                pass

    REPORTS_INDEX.write_text(json.dumps(index, indent=2), encoding="utf-8")
    return pdf_path


@app.route("/api/reports")
def api_reports():
    """List the last MAX_REPORTS PDF grade reports (FIFO queue)."""
    if not REPORTS_INDEX.exists():
        return jsonify(reports=[])
    try:
        index = json.loads(REPORTS_INDEX.read_text())
    except Exception:
        return jsonify(reports=[])
    # Only return entries whose file still exists
    valid = [e for e in index if (REPORTS_DIR / e["filename"]).exists()]
    return jsonify(reports=valid)


@app.route("/api/reports/<filename>")
def api_serve_report(filename):
    """Serve a PDF grade report for download."""
    if not re.match(r'^[\w\-]+\.pdf$', filename):
        return "Invalid filename", 400
    path = REPORTS_DIR / filename
    if not path.exists():
        return "Report not found", 404
    return send_file(str(path), mimetype="application/pdf",
                     as_attachment=True, download_name=filename)


@app.route("/api/grade-quiz-batch", methods=["POST"])
def api_grade_quiz_batch():
    """Grade all students in a quiz Excel file — returns updated xlsx with AI Score + AI Feedback columns."""
    import sys as _sys, io as _io, tempfile as _tempfile
    if str(SCRIPTS_DIR) not in _sys.path:
        _sys.path.insert(0, str(SCRIPTS_DIR))

    from grading.grade_submission import (
        SYSTEM_PROMPT, GRADING_PROVIDERS, DEFAULT_GRADING_MODELS, get_api_key,
        extract_rubric_criteria_from_docx, extract_rubric_criteria,
        ai_extract_rubric_criteria, DEFAULT_RUBRIC_CRITERIA,
        read_text_file, _snap_to_grade_band,
    )

    # ── inputs ──────────────────────────────────────────────────────────────
    xlsx_f = request.files.get("quiz_xlsx")
    if not xlsx_f or not xlsx_f.filename:
        return jsonify(success=False, error="No Excel file uploaded."), 400

    provider      = request.form.get("provider", "anthropic")
    model_req     = (request.form.get("model") or "").strip()
    quiz_question = (request.form.get("quiz_question") or "").strip()
    selected_rubric = (request.form.get("selected_rubric") or "").strip()

    if provider not in GRADING_PROVIDERS:
        return jsonify(success=False, error=f"Unknown provider: {provider}"), 400

    key_name, call_fn = GRADING_PROVIDERS[provider]
    api_key = get_api_key(key_name)
    if not api_key:
        return jsonify(success=False, error=f"{key_name.upper()}_API_KEY not set."), 400
    model = model_req or DEFAULT_GRADING_MODELS.get(provider, "")

    # ── resolve rubric ───────────────────────────────────────────────────────
    rubric_path: Path | None = None
    rubric_f = request.files.get("rubric")
    if rubric_f and rubric_f.filename:
        tmp = _tempfile.NamedTemporaryFile(delete=False, suffix=Path(rubric_f.filename).suffix)
        rubric_f.save(tmp.name)
        rubric_path = Path(tmp.name)
    elif selected_rubric:
        rubric_path = _resolve_selected_path(
            selected_rubric,
            allowed_roots=[LIBRARY_RUBRICS_DIR, DEFAULT_RUBRIC_DIR, LIBRARY_ASSIGNMENTS_DIR, LIBRARY_QUIZZES_DIR],
            allowed_exts=RUBRIC_ALLOWED_EXTS,
        )

    rubric_text = read_text_file(rubric_path) if rubric_path else ""
    rubric_criteria: list[dict] = []

    if rubric_path and rubric_path.suffix.lower() == ".json":
        try:
            data = json.loads(rubric_path.read_text(encoding="utf-8"))
            raw = data.get("criteria", [])
            rubric_criteria = [
                {"criterion_id": f"C{i+1}", "criterion_name": c.get("criterion_name", f"C{i+1}"),
                 "max_points": float(c.get("max_points", 0)), "checklist_items": c.get("checklist_items", [])}
                for i, c in enumerate(raw) if float(c.get("max_points", 0)) > 0
            ]
        except Exception:
            pass
    if not rubric_criteria and rubric_path and rubric_path.suffix.lower() == ".docx":
        rubric_criteria = extract_rubric_criteria_from_docx(rubric_path)
    if not rubric_criteria:
        rubric_criteria = extract_rubric_criteria(rubric_text)
    if not rubric_criteria and rubric_text:
        rubric_criteria = ai_extract_rubric_criteria(rubric_text, api_key, provider)
    if not rubric_criteria:
        rubric_criteria = list(DEFAULT_RUBRIC_CRITERIA)

    total_max    = sum(c["max_points"] for c in rubric_criteria)
    criteria_str = json.dumps(rubric_criteria, indent=2)

    # ── read Excel ───────────────────────────────────────────────────────────
    try:
        import openpyxl
        wb = openpyxl.load_workbook(xlsx_f)
        ws = wb.active
    except Exception as e:
        return jsonify(success=False, error=f"Failed to read Excel: {e}"), 400

    # Detect header columns (row 1)
    headers = {}
    for c in range(1, ws.max_column + 1):
        h = str(ws.cell(1, c).value or "").strip().lower()
        headers[h] = c

    def _find_col(*keywords):
        for kw in keywords:
            for h, c in headers.items():
                if kw in h:
                    return c
        return None

    student_col = _find_col("student number", "student id", "student num", "student", "id") or 1
    answer_col  = _find_col("student answer", "answer", "response", "submission") or 2

    # Add output columns
    score_col    = ws.max_column + 1
    feedback_col = score_col + 1
    ws.cell(1, score_col).value    = "AI Score"
    ws.cell(1, feedback_col).value = "AI Feedback"

    # Style header cells
    from openpyxl.styles import PatternFill, Font
    teal_fill = PatternFill("solid", fgColor="0891B2")
    white_bold = Font(bold=True, color="FFFFFF")
    for c in (score_col, feedback_col):
        ws.cell(1, c).fill = teal_fill
        ws.cell(1, c).font = white_bold

    # ── grade each student ───────────────────────────────────────────────────
    graded = 0
    errors = []
    for row in range(2, ws.max_row + 1):
        answer = str(ws.cell(row, answer_col).value or "").strip()
        if not answer:
            continue

        user_msg = (
            f"Grade the following student quiz answer using the rubric criteria.\n\n"
            f"RUBRIC CRITERIA (JSON):\n{criteria_str}\n\n"
            f"RUBRIC TEXT:\n{rubric_text or '(Use criteria above)'}\n\n"
            f"QUIZ QUESTION:\n{quiz_question or '(Grade based on rubric criteria above)'}\n\n"
            f"STUDENT ANSWER:\n{answer[:3000]}\n\n"
            f"Return standard JSON grading schema."
        )

        try:
            resp   = call_fn(model=model, api_key=api_key, system=SYSTEM_PROMPT, user=user_msg)
            result = resp.get("result", {})

            # Extract score from criterion_scores (most accurate) or overall_score fallback
            score    = 0.0
            feedback = ""
            if isinstance(result, dict):
                crit_raw = (result.get("criterion_scores") or result.get("criteria_scores")
                            or result.get("criteria") or [])
                if isinstance(crit_raw, list) and crit_raw:
                    for item in crit_raw:
                        if not isinstance(item, dict):
                            continue
                        awarded = float(item.get("awarded_points", 0) or 0)
                        pct     = float(item.get("checklist_pct", 0) or 0)
                        mp      = float(item.get("max_points", 0) or 0)
                        if not awarded and pct and mp:
                            awarded = _snap_to_grade_band(pct, mp)
                        score += awarded
                else:
                    score = float(result.get("overall_score", 0) or 0)
                feedback = str(result.get("overall_feedback", "") or "").strip()

            score = round(min(total_max, max(0.0, score)), 1)
            ws.cell(row, score_col).value    = score
            ws.cell(row, feedback_col).value = feedback
            graded += 1
        except Exception as exc:
            ws.cell(row, score_col).value    = "ERROR"
            ws.cell(row, feedback_col).value = str(exc)[:120]
            errors.append(f"Row {row}: {exc}")

    # ── return updated xlsx ──────────────────────────────────────────────────
    buf = _io.BytesIO()
    wb.save(buf)
    buf.seek(0)
    orig_stem  = re.sub(r"[^\w\-]", "_", Path(xlsx_f.filename).stem)[:40]
    out_name   = f"{orig_stem}_AI_Graded.xlsx"
    return send_file(
        buf,
        mimetype="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        as_attachment=True,
        download_name=out_name,
    )


if __name__ == "__main__":
    print("\n  AI Auto Grader — Web Interface")
    print("  http://localhost:5000\n")
    debug = str(os.getenv("FLASK_DEBUG", "0")).strip().lower() in {"1", "true", "yes"}
    host = os.getenv("FLASK_HOST", "127.0.0.1")
    port = int(os.getenv("FLASK_PORT", "5000"))
    app.run(debug=debug, host=host, port=port)
