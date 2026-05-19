#!/usr/bin/env python3
"""
Create readable DOCX files from per-file extraction JSON (PDF sources only).

For each *.pdf.extraction.json:
- Reconstruct content page-wise.
- Insert extracted image text where image chunks appear.
- Save DOCX under a mirrored folder structure.
"""

from __future__ import annotations

import argparse
import json
import re
from pathlib import Path
from typing import Any

try:
    from docx import Document
except Exception as exc:  # pragma: no cover
    raise SystemExit(
        "python-docx is required. Install with: "
        "pip install python-docx (or pip install -r requirements.txt)"
    ) from exc


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(description="Export PDF extraction JSON to DOCX.")
    parser.add_argument(
        "--per-file-json-dir",
        required=True,
        help="Root folder containing *.extraction.json files from pipeline output.",
    )
    parser.add_argument(
        "--output-dir",
        required=True,
        help="Folder where DOCX exports should be written.",
    )
    return parser.parse_args()


def safe_mkdir(path: Path) -> None:
    path.mkdir(parents=True, exist_ok=True)


def clean_text(text: str) -> str:
    text = text.replace("\x00", " ")
    text = text.replace("\u2014", "-")
    text = re.sub(r"\s+", " ", text).strip()
    return text


def image_text_from_chunk(chunk: dict[str, Any]) -> str:
    md = chunk.get("metadata", {})
    extracted = clean_text(md.get("extracted_text", "") or "")
    if extracted:
        return extracted

    visible = md.get("visible_text_lines")
    if isinstance(visible, list):
        lines = [clean_text(str(x)) for x in visible if clean_text(str(x))]
        if lines:
            return "\n".join(lines)

    content = chunk.get("content", "") or ""
    content = str(content)
    marker = "Extracted visible text:"
    if marker in content:
        after = content.split(marker, 1)[1]
        # Stop before the description section if present.
        if "\n\nDescription:" in after:
            after = after.split("\n\nDescription:", 1)[0]
        lines = []
        for raw in after.splitlines():
            t = raw.strip()
            if t.startswith("- "):
                t = t[2:]
            t = clean_text(t)
            if t and t != "[none]":
                lines.append(t)
        if lines:
            return "\n".join(lines)
    return clean_text(content) or "[no_image_text_extracted]"


def chunk_sort_key(chunk: dict[str, Any]) -> tuple[int, float, float, int, int]:
    md = chunk.get("metadata", {})
    page = int(md.get("page") or 0)
    ct = md.get("content_type", "")
    bbox = md.get("bbox")

    y0 = 1e9
    x0 = 1e9
    if isinstance(bbox, dict):
        by = bbox.get("y0")
        bx = bbox.get("x0")
        if isinstance(by, (int, float)):
            y0 = float(by)
        if isinstance(bx, (int, float)):
            x0 = float(bx)

    # Keep text slightly before image when bbox ties.
    type_rank = 0 if ct == "text" else 1

    block_id = md.get("block_id", "")
    block_num = 0
    if isinstance(block_id, str) and block_id.startswith("T"):
        try:
            block_num = int(block_id[1:])
        except Exception:
            block_num = 0

    chunk_idx = int(md.get("chunk_index_in_block") or md.get("image_index_on_page") or 0)
    return (page, y0, x0, type_rank, block_num * 10000 + chunk_idx)


def export_pdf_json_to_docx(json_path: Path, per_file_json_root: Path, output_dir: Path) -> Path | None:
    payload = json.loads(json_path.read_text(encoding="utf-8"))
    source_path = str(payload.get("source_path", ""))
    if not source_path.lower().endswith(".pdf"):
        return None

    chunks = payload.get("chunks", [])
    if not isinstance(chunks, list):
        chunks = []

    chunks = sorted(chunks, key=chunk_sort_key)

    rel = json_path.relative_to(per_file_json_root)
    # Remove trailing ".extraction.json" and convert to .docx.
    stem = str(rel)
    if stem.endswith(".extraction.json"):
        stem = stem[: -len(".extraction.json")]
    out_path = output_dir / f"{stem}.docx"
    safe_mkdir(out_path.parent)

    doc = Document()
    doc.add_heading(f"Extracted Reconstruction: {Path(source_path).name}", level=1)
    doc.add_paragraph(f"Source: {source_path}")

    current_page = None
    image_counter = 0
    for chunk in chunks:
        md = chunk.get("metadata", {})
        page = md.get("page")
        if isinstance(page, int) and page != current_page:
            current_page = page
            doc.add_heading(f"Page {page}", level=2)

        ctype = md.get("content_type", "")
        if ctype == "text":
            text = clean_text(str(chunk.get("content", "") or ""))
            if text:
                doc.add_paragraph(text)
        elif ctype == "image_description":
            image_counter += 1
            img_text = image_text_from_chunk(chunk)
            desc = clean_text(str(md.get("description", "") or ""))
            p = doc.add_paragraph()
            p.add_run(f"[Image {image_counter} extracted text]").bold = True
            p.add_run("\n" + img_text)
            if desc:
                p.add_run("\n\n[Image summary] ").bold = True
                p.add_run(desc)

    doc.save(out_path)
    return out_path


def main() -> int:
    args = parse_args()
    per_file_json_dir = Path(args.per_file_json_dir).expanduser().resolve()
    output_dir = Path(args.output_dir).expanduser().resolve()

    if not per_file_json_dir.exists() or not per_file_json_dir.is_dir():
        raise SystemExit(f"Invalid per-file-json-dir: {per_file_json_dir}")

    safe_mkdir(output_dir)

    files = sorted(per_file_json_dir.rglob("*.extraction.json"))
    written: list[Path] = []
    for fp in files:
        out = export_pdf_json_to_docx(fp, per_file_json_dir, output_dir)
        if out is not None:
            written.append(out)

    print(f"Input extraction files found: {len(files)}")
    print(f"PDF DOCX files generated: {len(written)}")
    if written:
        print("Sample outputs:")
        for p in written[:5]:
            print(str(p))
    return 0


if __name__ == "__main__":
    raise SystemExit(main())

