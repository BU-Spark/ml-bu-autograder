"""
blueprints/rubric.py — /api/generate-rubric
"""
from __future__ import annotations

import json
import os
import re
import sys
from pathlib import Path

from flask import Blueprint, jsonify, request

from web.config import LIBRARY_RUBRICS_DIR, SCRIPTS_DIR
from web.utils.files import _assignment_number, _display_name, _is_allowed_ext, _quiz_number

rubric_bp = Blueprint("rubric", __name__)


def _read_pdf_text(raw: bytes) -> str:
    try:
        import fitz
        doc = fitz.open(stream=raw, filetype="pdf")
        return "\n".join(page.get_text() for page in doc).strip()
    except Exception as exc:
        raise RuntimeError(f"Failed to read PDF: {exc}") from exc


def _read_docx_text(raw: bytes) -> str:
    import io as _io
    try:
        from docx import Document
        doc   = Document(_io.BytesIO(raw))
        parts = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
        for tbl in doc.tables:
            for row in tbl.rows:
                line = " | ".join(c.text.strip() for c in row.cells if c.text.strip())
                if line:
                    parts.append(line)
        return "\n".join(parts).strip()
    except Exception as exc:
        raise RuntimeError(f"Failed to read DOCX: {exc}") from exc


@rubric_bp.route("/api/generate-rubric", methods=["POST"])
def api_generate_rubric():
    if str(SCRIPTS_DIR) not in sys.path:
        sys.path.insert(0, str(SCRIPTS_DIR))

    api_key = os.getenv("ANTHROPIC_API_KEY")
    if not api_key:
        return jsonify(success=False, error="ANTHROPIC_API_KEY is not set."), 400

    # ── assignment text ──────────────────────────────────────────────────────
    assignment_text = (request.form.get("assignment_text") or "").strip()

    if not assignment_text and "assignment_file" in request.files:
        f = request.files["assignment_file"]
        if f and f.filename:
            if not _is_allowed_ext(f.filename, {".txt", ".pdf", ".md"}):
                return jsonify(success=False, error="Assignment file must be .txt, .pdf, or .md"), 400
            raw = f.read()
            try:
                assignment_text = _read_pdf_text(raw) if f.filename.lower().endswith(".pdf") \
                    else raw.decode("utf-8", errors="replace").strip()
            except RuntimeError as exc:
                return jsonify(success=False, error=str(exc)), 400

    if not assignment_text:
        lib_path_str = (request.form.get("assignment_library_path") or "").strip()
        if lib_path_str:
            lib_p = Path(lib_path_str)
            if not lib_p.exists():
                return jsonify(success=False, error=f"Library file not found: {lib_path_str}"), 400
            ext_lp = lib_p.suffix.lower()
            try:
                if ext_lp == ".pdf":
                    assignment_text = _read_pdf_text(lib_p.read_bytes())
                elif ext_lp == ".docx":
                    assignment_text = _read_docx_text(lib_p.read_bytes())
                else:
                    assignment_text = lib_p.read_text(encoding="utf-8", errors="replace").strip()
            except RuntimeError as exc:
                return jsonify(success=False, error=str(exc)), 400

    if not assignment_text:
        return jsonify(success=False,
                       error="Provide assignment_text, upload an assignment_file, or select a saved assignment."), 400

    # ── existing rubric (optional — triggers enhance mode) ───────────────────
    existing_rubric = (request.form.get("existing_rubric") or "").strip()
    if not existing_rubric and "existing_rubric_file" in request.files:
        rf = request.files["existing_rubric_file"]
        if rf and rf.filename:
            if not _is_allowed_ext(rf.filename, {".txt", ".md", ".pdf", ".docx", ".json"}):
                return jsonify(success=False, error="Rubric file must be .txt, .md, .pdf, .docx, or .json"), 400
            raw = rf.read()
            ext = Path(rf.filename).suffix.lower()
            try:
                if ext == ".pdf":
                    existing_rubric = _read_pdf_text(raw)
                elif ext == ".docx":
                    existing_rubric = _read_docx_text(raw)
                else:
                    existing_rubric = raw.decode("utf-8", errors="replace").strip()
            except RuntimeError as exc:
                return jsonify(success=False, error=str(exc)), 400

    instructions = (request.form.get("instructions") or "").strip()
    model        = (request.form.get("model") or "claude-sonnet-4-6").strip()

    try:
        from rubric_gen.generate_rubric import generate_rubric, enhance_rubric, rubric_to_dict
    except ImportError as exc:
        return jsonify(success=False, error=f"Could not import rubric_gen: {exc}"), 500

    try:
        if existing_rubric:
            rubric = enhance_rubric(assignment_text, existing_rubric,
                                    instructions=instructions, model=model, api_key=api_key)
            mode = "enhanced"
        else:
            rubric = generate_rubric(assignment_text, instructions=instructions,
                                     model=model, api_key=api_key)
            mode = "generated"

        rubric_dict  = rubric_to_dict(rubric)
        saved_path   = None
        saved_display = None
        save_name    = (request.form.get("save_name") or "").strip()
        if save_name:
            LIBRARY_RUBRICS_DIR.mkdir(parents=True, exist_ok=True)
            num  = _assignment_number(save_name)
            qnum = _quiz_number(save_name)
            if num:
                canonical = f"assignment_{num}_rubric.json"
            elif qnum:
                canonical = f"quiz_{qnum}_rubric.json"
            elif (m := re.search(r"\d+", save_name)):
                canonical = f"assignment_{m.group()}_rubric.json"
            else:
                safe      = re.sub(r"[^\w\-]", "_", save_name).strip("_") or "generated_rubric"
                canonical = f"{safe}.json"
            dest = LIBRARY_RUBRICS_DIR / canonical
            dest.write_text(json.dumps(rubric_dict, indent=2), encoding="utf-8")
            saved_path    = str(dest)
            saved_display = _display_name(canonical, kind="rubric")

        return jsonify(success=True, mode=mode, rubric=rubric_dict,
                       saved_path=saved_path, saved_display=saved_display)
    except Exception as exc:
        return jsonify(success=False, error=str(exc)), 500
