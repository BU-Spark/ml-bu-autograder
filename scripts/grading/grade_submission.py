#!/usr/bin/env python3
"""
grade_submission.py — Grade a student submission using retrieved lecture context.
"""
from __future__ import annotations

import argparse
import json
import re
import sys
from pathlib import Path
from typing import Any

# Allow `from core.config import ...` regardless of where the script is run from.
SCRIPTS_DIR = Path(__file__).resolve().parents[1]
if str(SCRIPTS_DIR) not in sys.path:
    sys.path.insert(0, str(SCRIPTS_DIR))

from core.config import get_api_key, load_environment


# ---------------------------------------------------------------------------
# I/O helpers
# ---------------------------------------------------------------------------

def read_jsonl(path: Path) -> list[dict[str, Any]]:
    rows: list[dict[str, Any]] = []
    with path.open("r", encoding="utf-8") as f:
        for line in f:
            line = line.strip()
            if line:
                rows.append(json.loads(line))
    return rows


def write_json(path: Path, obj: Any) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_text(json.dumps(obj, indent=2, ensure_ascii=False), encoding="utf-8")


def anonymize_label(text: str) -> str:
    """Remove quality hints such as 'good example' / 'bad example' from labels."""
    t = str(text or "")
    t = re.sub(r"(?i)\bgood\s+example\b", "example", t)
    t = re.sub(r"(?i)\bbad\s+example\b", "example", t)
    return re.sub(r"\s+", " ", t).strip()


def _try_decode_json_object(text: str) -> dict[str, Any] | None:
    decoder = json.JSONDecoder()
    s = str(text or "")
    for i, ch in enumerate(s):
        if ch != "{":
            continue
        try:
            obj, _ = decoder.raw_decode(s[i:])
        except json.JSONDecodeError:
            continue
        if isinstance(obj, dict):
            return obj
    return None


def _parse_json_from_text(raw: str) -> dict[str, Any]:
    text = str(raw or "").strip()
    if not text:
        return {"raw_response": "", "parse_error": True}

    try:
        parsed = json.loads(text)
        if isinstance(parsed, dict):
            return parsed
    except json.JSONDecodeError:
        pass

    # Common case: model returns ```json ... ``` wrapper.
    fence = re.search(r"```(?:json)?\s*(.*?)\s*```", text, flags=re.IGNORECASE | re.DOTALL)
    if fence:
        inner = fence.group(1).strip()
        try:
            parsed = json.loads(inner)
            if isinstance(parsed, dict):
                return parsed
        except json.JSONDecodeError:
            maybe = _try_decode_json_object(inner)
            if maybe is not None:
                return maybe

    maybe = _try_decode_json_object(text)
    if maybe is not None:
        return maybe

    return {"raw_response": text, "parse_error": True}


def _has_expected_grading_schema(result: dict[str, Any]) -> bool:
    if not isinstance(result, dict):
        return False
    return any(
        key in result
        for key in (
            "criterion_scores",
            "criterionScores",
            "criteria_scores",
            "criteria",
            "score_breakdown",
            "scoreBreakdown",
        )
    )


def _is_single_criterion_fragment(result: dict[str, Any]) -> bool:
    if not isinstance(result, dict):
        return False
    has_id = any(k in result for k in ("criterion_id", "criterionId"))
    has_points = any(k in result for k in ("awarded_points", "awardedPoints", "score", "points"))
    return has_id and has_points and not _has_expected_grading_schema(result)


# ---------------------------------------------------------------------------
# Data loading
# ---------------------------------------------------------------------------

def load_student_content(
    chunks_jsonl: Path,
    student_path_filter: str | None,
) -> tuple[str, list[dict[str, Any]]]:
    """
    Return (assembled_text, student_chunks) sorted by document_order.
    If student_path_filter is given, only include chunks whose source_path
    contains that substring (case-insensitive).
    """
    all_chunks = read_jsonl(chunks_jsonl)
    student_chunks = [
        c for c in all_chunks
        if str(c.get("metadata", {}).get("source_type", "")).lower() == "student"
    ]
    if student_path_filter:
        filt = student_path_filter.lower()
        student_chunks = [
            c for c in student_chunks
            if filt in str(c.get("metadata", {}).get("source_path", "")).lower()
        ]

    student_chunks.sort(
        key=lambda c: (
            c.get("metadata", {}).get("document_order", 9999),
            c.get("metadata", {}).get("sort_key", ""),
        )
    )

    parts: list[str] = []
    seen: set[str] = set()
    for c in student_chunks:
        text = str(c.get("content", "")).strip()
        if not text or text in seen:
            continue
        seen.add(text)

        # Label each chunk by content type so the grader knows what it's reading.
        meta = c.get("metadata", {}) or {}
        ctype = str(meta.get("content_type", "text")).lower()
        page = meta.get("page_number", "?")
        fmt = str(meta.get("format", "")).upper()

        if ctype == "image_description":
            image_type = meta.get("image_type") or "image"
            ocr_in_content = "OCR extracted text" in text
            label = f"[IMAGE | {fmt} Page {page} | type={image_type}{'| includes OCR' if ocr_in_content else ''}]"
        elif ctype == "table":
            tbl_idx = meta.get("table_index", "")
            label = f"[TABLE | {fmt} Page {page}{' #' + str(tbl_idx) if tbl_idx else ''}]"
        else:
            label = f"[TEXT | {fmt} Page {page}]"

        parts.append(f"{label}\n{text}")

    return "\n\n---\n\n".join(parts), student_chunks


def load_lecture_context(
    retrieval_jsonl: Path,
    student_path_filter: str | None,
    max_chars: int,
) -> str:
    """
    Collect unique lecture snippets from retrieval_results.jsonl,
    ordered by how often they appear (most-retrieved first).
    """
    rows = read_jsonl(retrieval_jsonl)
    if student_path_filter:
        filt = student_path_filter.lower()
        rows = [
            r for r in rows
            if filt in str(r.get("student_source_path", "")).lower()
        ]

    # Count how often each lecture document appears across all student chunks.
    doc_counts: dict[str, int] = {}
    for row in rows:
        results = row.get("results", {})
        docs = results.get("documents", [[]])[0]
        for doc in docs:
            doc = str(doc).strip()
            if doc:
                doc_counts[doc] = doc_counts.get(doc, 0) + 1

    # Sort most-retrieved first, then assemble up to max_chars.
    ranked = sorted(doc_counts.items(), key=lambda x: -x[1])
    parts: list[str] = []
    total = 0
    for doc, _ in ranked:
        if total + len(doc) > max_chars:
            break
        parts.append(doc)
        total += len(doc)

    return "\n\n---\n\n".join(parts)


# ---------------------------------------------------------------------------
# Prompt builder
# ---------------------------------------------------------------------------

DEFAULT_RUBRIC_CRITERIA: list[dict[str, Any]] = [
    # Generic fallback only — used when NO rubric file is supplied.
    # Always upload a rubric DOCX/PDF for assignment-specific and accurate scoring.
    # These generic criteria produce approximate grades and should not be used for submission.
    {"criterion_id": "C1", "criterion_name": "content_and_depth", "max_points": 40.0},
    {"criterion_id": "C2", "criterion_name": "structure_and_organization", "max_points": 30.0},
    {"criterion_id": "C3", "criterion_name": "requirements_alignment", "max_points": 20.0},
    {"criterion_id": "C4", "criterion_name": "clarity_and_writing", "max_points": 10.0},
]


SYSTEM_PROMPT = """\
You are an expert grader for a graduate-level Health Informatics course.

BLIND GRADING:
- Ignore filename/folder labels such as "good example" or "bad example".
- Grade only the submission content against the assignment instructions and rubric.

GRADING METHOD — FOLLOW EXACTLY FOR EACH CRITERION:

Step 1 — Identify checklist items:
  The rubric contains ☐ checkbox items under each criterion.
  Locate every ☐ item for the criterion being graded.
  If no ☐ items exist, use your holistic judgment to estimate checklist_pct.

Step 2 — Evaluate each ☐ item against the student submission:
  YES     = clearly demonstrated with evidence
  PARTIAL = partially addressed or implied
  NO      = missing, incorrect, or not mentioned

Step 3 — Compute checklist_pct:
  checklist_pct = (yes_count + 0.5 × partial_count) / total_items × 100

Step 4 — Apply GRADE BAND TABLE to get awarded_points:
  95–100% → multiply max_points by 1.000
  90–94%  → multiply max_points by 0.933
  85–89%  → multiply max_points by 0.900
  80–84%  → multiply max_points by 0.833
  75–79%  → multiply max_points by 0.800
  70–74%  → multiply max_points by 0.733
  65–69%  → multiply max_points by 0.700
  60–64%  → multiply max_points by 0.633
  55–59%  → multiply max_points by 0.600
  below 55% → multiply max_points by 0.500

  Round awarded_points to nearest 0.5.

IMPORTANT RULES:
- Images, diagrams, and tables in the submission ARE valid evidence — reference them as [IMAGE Page N] or [TABLE Page N].
- Do NOT deduct for items not explicitly in the rubric or assignment.
- Do NOT assign the same score to every criterion — differentiate based on evidence quality.
- missing_items must name specific ☐ checklist items that were NO or missing.

OUTPUT COMPACTNESS (CRITICAL — prevent truncation):
- justification: max 30 words
- evidence_refs: max 1 short quote or image/table reference
- missing_items: max 2 short phrases
- strengths/gaps/action_items: max 2 items, short phrases only
- No markdown, no text outside the JSON

Return ONLY valid JSON:
{
  "student_file": "<filename>",
  "criterion_scores": [
    {
      "criterion_id": "C1",
      "checklist_pct": <0-100 float>,
      "awarded_points": <number from grade band table>,
      "justification": "<max 30 words>",
      "evidence_refs": ["<short direct quote or [IMAGE Page N] reference>"],
      "missing_items": ["<specific missing checklist item>"]
    }
  ],
  "section_coverage": [
    {
      "section_id": "Q1",
      "status": "addressed|partial|missing",
      "evidence_refs": ["<short quote>"]
    }
  ],
  "overall_feedback": "<2-4 sentences>",
  "strengths": ["<bullet>"],
  "gaps": ["<bullet>"],
  "action_items": ["<specific improvement>"],
  "confidence": <0-1 float>
}
"""


def detect_expected_sections(text: str) -> list[str]:
    """
    Detect numbered sections like:
      1. ...       Q1. ...
      2) ...       Q2. ...
      Section 3 ...
    Returns stable IDs: ["Q1", "Q2", ...]
    """
    # Match both plain numbers (1. 2)) and Q-prefixed (Q1. Q2.)
    matches = re.findall(
        r"(?im)^\s*(?:section\s+)?(?:Q)?([1-9][0-9]?)\s*[\.\)]\s+",
        text or "",
    )
    ordered_unique: list[int] = []
    seen: set[int] = set()
    for m in matches:
        n = int(m)
        if n not in seen:
            seen.add(n)
            ordered_unique.append(n)
    ordered_unique.sort()
    return [f"Q{n}" for n in ordered_unique]


def _clean_ws(text: str) -> str:
    return re.sub(r"\s+", " ", str(text or "")).strip()


def _slugify(text: str) -> str:
    s = re.sub(r"[^a-zA-Z0-9]+", "_", _clean_ws(text).lower()).strip("_")
    return s or "criterion"


def extract_rubric_criteria(rubric_text: str) -> list[dict[str, Any]]:
    """
    Parse criteria from rubric text, looking for lines containing "(N points)".
    Returns a list like:
      [{"criterion_id":"C1","criterion_name":"...","max_points":30.0}, ...]
    """
    if not rubric_text:
        return []

    lines = [_clean_ws(line) for line in str(rubric_text).splitlines()]
    lines = [line for line in lines if line]

    parsed: list[dict[str, Any]] = []
    seen: set[tuple[str, int]] = set()

    def _is_grade_or_noise_line(s: str) -> bool:
        low = s.lower()
        if not low:
            return True
        if "feedback" in low and len(low) < 32:
            return True
        if "other feedback" in low:
            return True
        if "learning objective" in low:
            return True
        # "☐Other: _____" placeholder form fields (with or without checkbox glyph)
        if re.match(r"^\W*other\s*[:.]?\s*_+\s*$", low):
            return True
        if re.search(r"\b\d{2}\s*-\s*\d{2}\b", low):
            return True
        if re.search(r"\b\d{2}\s*-\s*100\b", low):
            return True
        if "pts" in low and re.search(r"\d", low):
            return True
        if re.fullmatch(r"[_\-\s|]+", s):
            return True
        return False

    for idx, line in enumerate(lines):
        for m in re.finditer(r"\((\d{1,3})\s*points?\)", line, flags=re.IGNORECASE):
            pts = int(m.group(1))
            # Match the DOCX parser's upper bound (500) so single-criterion rubrics
            # worth >100 pts are not silently dropped by the text parser.
            if pts <= 0 or pts > 500:
                continue

            left = _clean_ws(line[: m.start()])
            left = left.split("|")[0].strip(" |-:")
            # If points are on a dedicated line, pull criterion text from lines above.
            if len(left) < 8:
                collected: list[str] = []
                for j in range(idx - 1, max(-1, idx - 8), -1):
                    prev = _clean_ws(lines[j])
                    if not prev:
                        continue
                    # Checkbox lines (☐ / ☑) are feedback items from the PREVIOUS
                    # criterion; stop walking when we hit one.
                    if re.match(r"^\s*[☐☑]", prev):
                        break
                    if _is_grade_or_noise_line(prev):
                        continue
                    # stop at very broad section dividers
                    if re.search(r"(?i)\b(total|max\s+\d+\s+points?)\b", prev):
                        continue
                    collected.append(prev.split("|")[0].strip(" |-:"))
                    if len(collected) >= 4:
                        break
                if collected:
                    collected.reverse()
                    left = _clean_ws(" ".join(collected))

            if not left:
                continue
            if re.fullmatch(r"(?i)total", left):
                continue
            if re.search(r"(?i)\bgrades?\s+achieved\b", left):
                continue
            if re.search(r"(?i)\blearning\s+objective\b", left):
                continue

            key = (_slugify(left), pts)
            if key in seen:
                continue
            seen.add(key)

            parsed.append(
                {
                    "criterion_id": f"C{len(parsed) + 1}",
                    "criterion_name": left,
                    "max_points": float(pts),
                }
            )
            # Do NOT break — a line can have multiple criteria (e.g. "A (30 pts) or B (20 pts)")

    if not parsed:
        return []
    total = sum(float(c["max_points"]) for c in parsed)
    if total <= 0 or total > 1000:
        print(
            f"WARNING: Rubric total {total} pts is outside expected range (0–1000). "
            "Falling back to default criteria. Check rubric formatting.",
            file=sys.stderr,
        )
        return []
    return parsed


def _is_rubric_noise_line(s: str) -> bool:
    """Return True if a line is rubric boilerplate, not a real checklist item."""
    low = s.lower().strip()
    if not low:
        return True
    # Placeholder "Other: ____" lines
    if re.match(r"(?i)^other\s*[:.]?\s*_+\s*$", low):
        return True
    if re.match(r"(?i)^other\s+feedback", low):
        return True
    # Grade band lines like "95-100 (30 pts)"
    if re.search(r"\d{2}\s*[-–]\s*\d{2,3}", s):
        return True
    # Very short or all-punctuation
    if len(low) < 5:
        return True
    return False


def _extract_checklist_items_from_text(text: str) -> list[str]:
    """Extract checklist items from rubric cell text.

    Handles three formats:
    1. Lines starting with ☐/□/•/▪ checkbox characters
    2. Plain lines after a "Feedback:" header (common in CS rubric DOCX)
    3. Plain lines with no special prefix (all lines are criteria)
    """
    items: list[str] = []

    # Format 1: ☐/□/•/▪ prefixed lines
    CHECKBOX_CHARS = "☐□☑☒▪•◦"
    for line in text.splitlines():
        stripped = line.strip()
        if stripped and stripped[0] in CHECKBOX_CHARS:
            item = _clean_ws(stripped[1:].lstrip())
            if item and not _is_rubric_noise_line(item):
                items.append(item)

    if items:
        return items

    # Format 2: plain text lines after "Feedback:" header
    in_feedback = False
    found_feedback_header = False
    for line in text.splitlines():
        stripped = line.strip()
        if not stripped:
            continue
        if re.match(r"(?i)^feedback\s*[:.]?\s*$", stripped):
            in_feedback = True
            found_feedback_header = True
            continue
        if in_feedback and not _is_rubric_noise_line(stripped) and 5 <= len(stripped) <= 300:
            items.append(_clean_ws(stripped))

    if found_feedback_header:
        return items

    # Format 3: all non-noise lines are criteria (no header, no ☐ prefix)
    for line in text.splitlines():
        stripped = line.strip()
        if stripped and not _is_rubric_noise_line(stripped) and 8 <= len(stripped) <= 300:
            items.append(_clean_ws(stripped))

    return items


def extract_rubric_criteria_from_docx(path: Path) -> list[dict[str, Any]]:
    try:
        from docx import Document
    except Exception:
        return []

    if not path.exists() or path.suffix.lower() != ".docx":
        return []

    def _normalize_name(name: str) -> str:
        # Use only the first non-empty line as the criterion name
        first_line = ""
        for line in name.splitlines():
            stripped = line.strip()
            if stripped and not re.match(r"(?i)^\(?\d{1,3}\s*points?\)?", stripped):
                first_line = stripped
                break
        s = _clean_ws(first_line or name)
        # Remove repeated placeholder lines and trailing separators.
        s = re.sub(r"(?i)\bother:\s*_+\b", "", s)
        s = re.sub(r"(?i)\(\s*\d{1,3}\s*points?\s*\)", "", s)
        s = re.sub(r"_+", " ", s)
        s = re.sub(r"\s{2,}", " ", s).strip(" |:-")
        return s

    criteria: list[dict[str, Any]] = []
    seen: set[tuple[str, int]] = set()

    try:
        doc = Document(str(path))
    except Exception:
        return []

    for table in doc.tables:
        current_criterion_idx: int | None = None
        for row in table.rows:
            # Use raw cell text (preserving newlines) for checklist extraction
            raw_cells = [c.text for c in row.cells]
            cells = [_clean_ws(t) for t in raw_cells]
            if not any(cells):
                continue

            name_cell = cells[0] if cells else ""
            if not name_cell:
                continue
            low_name = name_cell.lower()
            if "learning objective" in low_name:
                continue
            if re.search(r"(?i)\btotal\b", name_cell):
                continue

            # Extract points from any cell in the row.
            point_candidates: list[int] = []
            for cell_text in cells:
                for m in re.finditer(r"\((\d{1,3})\s*points?\)", cell_text, flags=re.IGNORECASE):
                    point_candidates.append(int(m.group(1)))
                for m in re.finditer(r"\((\d{1,3})\s*pts?\)", cell_text, flags=re.IGNORECASE):
                    point_candidates.append(int(m.group(1)))

            if point_candidates:
                max_points = max([p for p in point_candidates if 0 < p <= 500], default=0)
                if max_points <= 0:
                    continue

                cname = _normalize_name(raw_cells[0] if raw_cells else name_cell)
                if len(cname) < 6:
                    continue
                key = (_slugify(cname), max_points)
                if key in seen:
                    continue
                seen.add(key)

                # Extract checklist items: try each column from longest to shortest
                # (longest text column most likely contains the criteria/feedback).
                # Fall back to scanning all cells if nothing found in any single column.
                checklist_items: list[str] = []
                col_candidates = sorted(raw_cells, key=len, reverse=True)
                for col_text in col_candidates:
                    checklist_items = _extract_checklist_items_from_text(col_text)
                    if checklist_items:
                        break
                if not checklist_items:
                    full_row_raw = "\n".join(raw_cells)
                    checklist_items = _extract_checklist_items_from_text(full_row_raw)

                criteria.append(
                    {
                        "criterion_id": f"C{len(criteria) + 1}",
                        "criterion_name": cname,
                        "max_points": float(max_points),
                        "checklist_items": checklist_items,
                    }
                )
                current_criterion_idx = len(criteria) - 1

            elif current_criterion_idx is not None:
                # Non-criterion row (no points pattern) — may contain ☐ items for the current criterion
                full_row_raw = "\n".join(raw_cells)
                new_items = _extract_checklist_items_from_text(full_row_raw)
                if new_items:
                    existing = criteria[current_criterion_idx].get("checklist_items", [])
                    # Deduplicate
                    existing_set = set(existing)
                    for item in new_items:
                        if item not in existing_set:
                            existing.append(item)
                            existing_set.add(item)
                    criteria[current_criterion_idx]["checklist_items"] = existing

    total = sum(float(c["max_points"]) for c in criteria)
    if not criteria or total <= 0 or total > 1000:
        if criteria:
            print(
                f"WARNING: Rubric table total {total} pts is outside expected range (0–1000). "
                "Falling back to default criteria. Check rubric formatting.",
                file=sys.stderr,
            )
        return []
    return criteria


def ai_extract_rubric_criteria(
    rubric_text: str,
    api_key: str,
    provider: str = "anthropic",
) -> list[dict[str, Any]]:
    """Use a fast LLM to extract criteria from a free-form rubric (haiku / gpt-4o-mini).

    Called as a fallback when the regex-based parsers fail (prose rubrics).
    """
    if not rubric_text or not api_key:
        return []

    RUBRIC_SYSTEM = (
        "You are a rubric parser. Read the rubric and return ONLY a valid JSON array — "
        "no markdown fences, no explanation. Each element must have exactly these keys:\n"
        '  "criterion_id": "C1", "C2", etc.\n'
        '  "criterion_name": short snake_case name\n'
        '  "max_points": number (integer or float)\n'
        '  "checklist_items": list of strings — specific requirements to check for\n'
        "Rules: extract ALL criteria with their EXACT point values from the rubric. "
        "Do NOT invent criteria. Make sure all max_points sum to the rubric total."
    )
    if len(rubric_text) > 6000:
        print(
            f"WARNING: AI rubric extractor truncating rubric from {len(rubric_text)} to 6000 chars — "
            "later criteria may be missed."
        )
    prompt = f"Extract all grading criteria from this rubric:\n\n{rubric_text[:6000]}"

    raw = ""
    try:
        if provider == "anthropic":
            from anthropic import Anthropic as _Anthropic
            _client = _Anthropic(api_key=api_key)
            _msg = _client.messages.create(
                model="claude-haiku-4-5-20251001",
                max_tokens=1024,
                temperature=0,
                system=RUBRIC_SYSTEM,
                messages=[{"role": "user", "content": prompt}],
            )
            raw = _msg.content[0].text.strip()
        elif provider == "openai":
            import openai as _openai
            _client = _openai.OpenAI(api_key=api_key)
            _resp = _client.chat.completions.create(
                model="gpt-4o-mini",
                max_tokens=1024,
                temperature=0,
                messages=[
                    {"role": "system", "content": RUBRIC_SYSTEM},
                    {"role": "user", "content": prompt},
                ],
            )
            raw = _resp.choices[0].message.content.strip()
        else:
            print(
                f"WARNING: ai_extract_rubric_criteria does not support provider '{provider}' — "
                "returning no criteria. Supported providers: anthropic, openai."
            )
            return []
    except Exception as exc:
        print(f"WARNING: AI rubric parsing failed ({provider}): {exc}")
        return []

    try:
        cleaned = re.sub(r"^```[a-z]*\n?", "", raw, flags=re.IGNORECASE)
        cleaned = re.sub(r"\n?```$", "", cleaned, flags=re.IGNORECASE).strip()
        data = json.loads(cleaned)
        if isinstance(data, dict):
            data = data.get("criteria", data.get("rubric_criteria", []))
        if not isinstance(data, list):
            return []
        criteria: list[dict[str, Any]] = []
        for i, c in enumerate(data):
            if not isinstance(c, dict):
                continue
            try:
                pts = float(c.get("max_points", 0))
            except Exception:
                continue
            if pts <= 0:
                continue
            name = str(c.get("criterion_name", f"criterion_{i+1}")).strip()
            criteria.append({
                "criterion_id": str(c.get("criterion_id", f"C{i+1}")),
                "criterion_name": name,
                "max_points": pts,
                "checklist_items": [str(x) for x in (c.get("checklist_items") or []) if x],
            })
        total = sum(c["max_points"] for c in criteria)
        if not criteria or total <= 0:
            return []
        if total > 1000:
            print(
                f"WARNING: AI rubric extraction produced implausible total ({total:.0f} pts > 1000) — "
                "discarding and falling back."
            )
            return []
        print(f"AI rubric parsing ({provider}): {len(criteria)} criteria, {total:.0f} total pts")
        return criteria
    except Exception as exc:
        print(f"WARNING: Failed to parse AI rubric extraction response: {exc}\nRaw: {raw[:300]}")
        return []


def flatten_system_blocks(blocks: list[dict[str, Any]]) -> str:
    """Collapse text-type system blocks into a plain string for OpenAI/Gemini providers.

    Only blocks with ``type == "text"`` are included. Any other type, or a missing
    type field, raises ValueError so callers cannot silently lose system content.
    """
    parts = []
    for b in blocks:
        btype = b.get("type")
        if btype == "text":
            parts.append(b["text"])
        else:
            raise ValueError(
                f"flatten_system_blocks: block missing 'type' or has unsupported type '{btype}' "
                "for non-Anthropic provider"
            )
    return "\n\n".join(parts)


def build_system_blocks(
    *,
    rubric_text: str,
    rubric_criteria: list[dict[str, Any]],
    assignment_text: str,
    expected_sections: list[str],
    enable_cache: bool = False,
) -> list[dict[str, Any]]:
    """Build structured system content blocks with optional Anthropic prompt caching.

    When ``enable_cache=True`` (Anthropic only), a single ``cache_control`` breakpoint is
    placed on block2 — the outermost stable boundary. This caches block1+block2 as one
    prefix. Block1 has no breakpoint to avoid a redundant intermediate cache write.
    """
    block1: dict[str, Any] = {"type": "text", "text": SYSTEM_PROMPT}

    if expected_sections:
        sections_line = f"Expected sections: {', '.join(expected_sections)}."
    elif assignment_text:
        sections_line = "Infer expected sections from assignment instructions."
    else:
        sections_line = "No assignment instructions provided; infer expected sections from submission."

    context_text = (
        "=== ASSIGNMENT INSTRUCTIONS (grade these questions) ===\n"
        f"{assignment_text or '(No assignment instructions provided)'}\n\n"
        "=== EXPECTED SECTIONS ===\n"
        f"{sections_line}\n\n"
        "=== RUBRIC CRITERIA (authoritative scoring dimensions) ===\n"
        f"{json.dumps(rubric_criteria, ensure_ascii=True, indent=2)}\n\n"
        "=== RUBRIC ===\n"
        f"{rubric_text or '(No rubric provided)'}"
    )

    block2: dict[str, Any] = {"type": "text", "text": context_text}
    if enable_cache:
        block2["cache_control"] = {"type": "ephemeral"}

    return [block1, block2]


def build_user_message(
    student_text: str,
    lecture_context: str,
    student_file: str,
    max_student_chars: int,
) -> str:
    student_excerpt = student_text[:max_student_chars]
    if len(student_text) > max_student_chars:
        student_excerpt += "\n\n[... content truncated ...]"

    return (
        f"STUDENT FILE: {student_file}\n\n"
        "=== LECTURE CONTEXT ===\n"
        f"{lecture_context}\n\n"
        "=== STUDENT SUBMISSION ===\n"
        f"{student_excerpt}"
    )


def read_text_file(path: Path | None) -> str:
    if path is None:
        return ""
    if not path.exists() or not path.is_file():
        return ""
    suffix = path.suffix.lower()

    # Plain text-like files
    if suffix in {".txt", ".md", ".csv", ".json"}:
        return path.read_text(encoding="utf-8", errors="ignore").strip()

    # DOCX
    if suffix == ".docx":
        try:
            from docx import Document
        except Exception:
            return ""
        try:
            doc = Document(str(path))
            parts: list[str] = []
            for p in doc.paragraphs:
                t = (p.text or "").strip()
                if t:
                    parts.append(t)
            # Also capture table text (important for rubrics)
            for table in doc.tables:
                for row in table.rows:
                    cells = [((c.text or "").strip()) for c in row.cells]
                    line = " | ".join([c for c in cells if c])
                    if line:
                        parts.append(line)
            return "\n".join(parts).strip()
        except Exception:
            return ""

    # PDF — PyMuPDF handles tables/multi-column better than pypdf; pypdf is the fallback.
    if suffix == ".pdf":
        try:
            import fitz  # PyMuPDF
            with fitz.open(str(path)) as doc:
                parts = []
                for page in doc:
                    t = (page.get_text() or "").strip()
                    if t:
                        parts.append(t)
            return "\n\n".join(parts).strip()
        except Exception as exc:
            print(f"INFO: PyMuPDF could not parse {path.name} ({exc}); falling back to pypdf.")
        try:
            from pypdf import PdfReader
            reader = PdfReader(str(path))
            parts = []
            for page in reader.pages:
                t = (page.extract_text() or "").strip()
                if t:
                    parts.append(t)
            return "\n\n".join(parts).strip()
        except Exception as exc:
            print(f"WARNING: pypdf also failed on {path.name}: {exc}")
            return ""

    # Fallback binary/text read
    return path.read_text(encoding="utf-8", errors="ignore").strip()


# ---------------------------------------------------------------------------
# LLM API calls (OpenAI, Gemini, Anthropic)
# ---------------------------------------------------------------------------

def call_openai(
    *,
    model: str,
    api_key: str,
    system: str | list[dict[str, Any]],
    user: str,
) -> dict[str, Any]:
    try:
        import openai
    except ImportError as exc:
        raise RuntimeError("openai package not installed. Run: pip install openai") from exc

    if isinstance(system, list):
        system = flatten_system_blocks(system)

    client = openai.OpenAI(api_key=api_key)
    response = client.chat.completions.create(
        model=model,
        max_tokens=4096,
        temperature=0,
        messages=[
            {"role": "system", "content": system},
            {"role": "user", "content": user},
        ],
        response_format={"type": "json_object"},
    )
    raw = response.choices[0].message.content.strip()
    parsed = _parse_json_from_text(raw)

    return {
        "result": parsed,
        "usage": {
            "input_tokens": response.usage.prompt_tokens,
            "output_tokens": response.usage.completion_tokens,
        },
        "model": response.model,
    }


def _dedupe_keep_order(values: list[str]) -> list[str]:
    out: list[str] = []
    seen: set[str] = set()
    for value in values:
        v = str(value or "").strip()
        if not v or v in seen:
            continue
        seen.add(v)
        out.append(v)
    return out


def _gemini_model_candidates(model: str) -> list[str]:
    requested = str(model or "").strip()
    if requested.startswith("models/"):
        requested = requested.split("/", 1)[1]
    if not requested:
        requested = "gemini-2.5-flash"

    candidates: list[str] = [requested]

    # Common alias handling.
    if requested.endswith("-flash") and not requested.endswith("-001"):
        candidates.append(f"{requested}-001")
    if requested.endswith("-pro") and not requested.endswith("-001"):
        candidates.append(f"{requested}-001")
    if requested.endswith("-001"):
        candidates.append(requested[:-4])

    # Known cross-version aliases to keep grading resilient across API key entitlements.
    if requested in {"gemini-1.5-flash", "gemini-1.5-flash-001", "gemini-1.5-flash-002"}:
        candidates.extend(["gemini-2.5-flash", "gemini-2.0-flash-001", "gemini-2.0-flash"])
    if requested == "gemini-2.0-flash":
        candidates.extend(["gemini-2.0-flash-001", "gemini-2.5-flash", "gemini-1.5-flash-002", "gemini-1.5-flash"])
    if requested == "gemini-2.0-flash-001":
        candidates.extend(["gemini-2.0-flash", "gemini-2.5-flash"])
    if requested == "gemini-2.0-flash-exp":
        candidates.extend(["gemini-2.0-flash-001", "gemini-2.5-flash", "gemini-1.5-flash-002", "gemini-1.5-flash"])
    if requested == "gemini-2.5-flash":
        candidates.extend(["gemini-2.0-flash-001", "gemini-2.0-flash"])

    return _dedupe_keep_order(candidates)


def _gemini_list_models(api_key: str) -> list[str]:
    import urllib.error
    import urllib.parse
    import urllib.request

    out: list[str] = []
    for api_version in ("v1beta", "v1"):
        endpoint = (
            f"https://generativelanguage.googleapis.com/{api_version}/models"
            f"?key={urllib.parse.quote(api_key, safe='')}"
        )
        req = urllib.request.Request(endpoint, method="GET")
        try:
            with urllib.request.urlopen(req, timeout=45) as resp:
                raw = resp.read().decode("utf-8", errors="ignore")
            parsed = json.loads(raw) if raw else {}
            for item in parsed.get("models", []) or []:
                name = str(item.get("name", "")).strip()
                if name.startswith("models/"):
                    name = name.split("/", 1)[1]
                if name:
                    out.append(name)
        except Exception:
            continue
    return _dedupe_keep_order(out)


def call_gemini(
    *,
    model: str,
    api_key: str,
    system: str | list[dict[str, Any]],
    user: str,
) -> dict[str, Any]:
    import urllib.error
    import urllib.parse
    import urllib.request

    if isinstance(system, list):
        system = flatten_system_blocks(system)

    model_candidates = _gemini_model_candidates(model)
    available_models = _gemini_list_models(api_key)
    if available_models:
        preferred_available = [
            m
            for m in available_models
            if (
                m in {"gemini-2.5-flash", "gemini-2.0-flash", "gemini-2.0-flash-001", "gemini-2.5-pro"}
                or (m.startswith("gemini-") and ("flash" in m or "pro" in m) and "tts" not in m and "preview" not in m)
            )
        ]
        model_candidates = _dedupe_keep_order(model_candidates + preferred_available)
    # Use a broadly compatible payload shape across Gemini endpoint/model variants.
    # Some variants reject system_instruction / responseMimeType.
    body: dict[str, Any] = {
        "contents": [
            {
                "parts": [{"text": f"{system}\n\n{user}"}],
            }
        ],
        "generationConfig": {
            "temperature": 0,
            "maxOutputTokens": 4096,
        },
    }
    parsed_response: dict[str, Any] | None = None
    chosen_model = model_candidates[0]
    chosen_api_version: str | None = None
    attempts: list[str] = []

    for candidate in model_candidates:
        for api_version in ("v1beta", "v1"):
            endpoint = (
                f"https://generativelanguage.googleapis.com/{api_version}/models/"
                f"{urllib.parse.quote(candidate, safe='')}:generateContent"
                f"?key={urllib.parse.quote(api_key, safe='')}"
            )
            req = urllib.request.Request(
                endpoint,
                data=json.dumps(body).encode("utf-8"),
                headers={"Content-Type": "application/json"},
                method="POST",
            )
            try:
                with urllib.request.urlopen(req, timeout=120) as resp:
                    raw_response = resp.read().decode("utf-8", errors="ignore")
                parsed_response = json.loads(raw_response) if raw_response else {}
                chosen_model = candidate
                chosen_api_version = api_version
                break
            except urllib.error.HTTPError as e:
                error_body = e.read().decode("utf-8", errors="ignore")
                attempts.append(f"{api_version}/{candidate}: HTTP {e.code}")
                # 404 => model/endpoint mismatch candidate, continue fallback.
                # 400 can also be model-variant incompatibility; continue trying candidates.
                if e.code not in {400, 404}:
                    raise RuntimeError(
                        f"Gemini API error {e.code} for {api_version}/{candidate}: {error_body}"
                    ) from e
                continue
        if parsed_response is not None:
            break

    if parsed_response is None:
        available = _gemini_list_models(api_key)
        available_preview = ", ".join(available[:12]) if available else "none returned by API"
        tried = ", ".join(attempts) if attempts else "no attempts"
        raise RuntimeError(
            "Gemini model not found for this API key.\n"
            f"Requested: {model}\n"
            f"Tried: {tried}\n"
            f"Available models (sample): {available_preview}"
        )

    # Extract text from Gemini response
    candidates = parsed_response.get("candidates", [])
    if not candidates:
        return {
            "result": {"raw_response": json.dumps(parsed_response), "parse_error": True},
            "usage": {"input_tokens": 0, "output_tokens": 0},
            "model": chosen_model,
        }

    text_parts: list[str] = []
    for cand in candidates:
        parts = cand.get("content", {}).get("parts", [])
        for part in parts:
            text_val = str(part.get("text", "")).strip()
            if text_val:
                text_parts.append(text_val)
    raw_text = "\n".join(text_parts).strip()
    parsed = _parse_json_from_text(raw_text)

    # If parsing failed or schema is incomplete (often truncated output), do one compact retry.
    needs_retry = (
        isinstance(parsed, dict)
        and (
            bool(parsed.get("parse_error"))
            or _is_single_criterion_fragment(parsed)
            or not _has_expected_grading_schema(parsed)
        )
    )
    if needs_retry and chosen_api_version:
        retry_suffix = (
            "\n\nIMPORTANT: Previous output was not valid JSON or was truncated.\n"
            "Return ONLY one valid compact JSON object, no markdown.\n"
            "Required top-level keys: student_file, criterion_scores, section_coverage, "
            "overall_feedback, strengths, gaps, action_items, confidence.\n"
            "criterion_scores MUST include one entry per rubric criterion.\n"
            "Keep each criterion justification <= 20 words.\n"
            "Keep evidence_refs to max 1 item per criterion.\n"
            "Keep missing_items to max 1 item per criterion.\n"
            "Keep strengths/gaps/action_items to max 2 items each."
        )
        retry_body: dict[str, Any] = {
            "contents": [
                {
                    "parts": [{"text": f"{system}\n\n{user}{retry_suffix}"}],
                }
            ],
            "generationConfig": {
                "temperature": 0,
                "maxOutputTokens": 4096,
            },
        }
        retry_endpoint = (
            f"https://generativelanguage.googleapis.com/{chosen_api_version}/models/"
            f"{urllib.parse.quote(chosen_model, safe='')}:generateContent"
            f"?key={urllib.parse.quote(api_key, safe='')}"
        )
        retry_req = urllib.request.Request(
            retry_endpoint,
            data=json.dumps(retry_body).encode("utf-8"),
            headers={"Content-Type": "application/json"},
            method="POST",
        )
        try:
            with urllib.request.urlopen(retry_req, timeout=120) as resp:
                retry_raw_response = resp.read().decode("utf-8", errors="ignore")
            retry_parsed_response = json.loads(retry_raw_response) if retry_raw_response else {}
            retry_candidates = retry_parsed_response.get("candidates", [])
            retry_text_parts: list[str] = []
            for cand in retry_candidates:
                parts = cand.get("content", {}).get("parts", [])
                for part in parts:
                    text_val = str(part.get("text", "")).strip()
                    if text_val:
                        retry_text_parts.append(text_val)
            retry_raw_text = "\n".join(retry_text_parts).strip()
            retry_parsed = _parse_json_from_text(retry_raw_text)
            if isinstance(retry_parsed, dict) and not retry_parsed.get("parse_error"):
                parsed = retry_parsed
                parsed_response = retry_parsed_response
        except Exception:
            # Keep original parse result if retry fails.
            pass

    usage_obj = parsed_response.get("usageMetadata", {}) or {}
    return {
        "result": parsed,
        "usage": {
            "input_tokens": int(usage_obj.get("promptTokenCount", 0) or 0),
            "output_tokens": int(usage_obj.get("candidatesTokenCount", 0) or 0),
        },
        "model": chosen_model,
    }


def call_anthropic(
    *,
    model: str,
    api_key: str,
    system: str | list[dict[str, Any]],
    user: str,
) -> dict[str, Any]:
    """Call the Anthropic Messages API and return ``{"result", "usage", "model"}``.

    The ``system`` parameter accepts either a plain string or a list of content blocks.
    Passing blocks with ``cache_control`` enables Anthropic prompt caching.
    """
    try:
        from anthropic import Anthropic
    except ImportError as exc:
        raise RuntimeError("anthropic package not installed. Run: pip install anthropic") from exc

    client = Anthropic(api_key=api_key)
    message = client.messages.create(
        model=model,
        max_tokens=4096,
        temperature=0,
        system=system,
        messages=[
            {"role": "user", "content": user},
        ],
    )
    raw = message.content[0].text.strip()
    parsed = _parse_json_from_text(raw)

    usage_obj = getattr(message, "usage", None)
    return {
        "result": parsed,
        "usage": {
            "input_tokens": int(getattr(usage_obj, "input_tokens", 0) or 0),
            "output_tokens": int(getattr(usage_obj, "output_tokens", 0) or 0),
            "cache_creation_input_tokens": int(getattr(usage_obj, "cache_creation_input_tokens", 0) or 0),
            "cache_read_input_tokens": int(getattr(usage_obj, "cache_read_input_tokens", 0) or 0),
        },
        "model": message.model,
    }


# Provider → caller mapping
GRADING_PROVIDERS = {
    "openai": ("openai", call_openai),
    "gemini": ("gemini", call_gemini),
    "anthropic": ("anthropic", call_anthropic),
}

# Default models per provider
DEFAULT_GRADING_MODELS = {
    "openai": "gpt-4o-2024-11-20",
    "gemini": "gemini-2.5-flash",
    "anthropic": "claude-sonnet-4-20250514",
}


def _to_float(value: Any) -> float | None:
    try:
        return float(value)
    except Exception:
        return None


def _to_string_list(value: Any) -> list[str]:
    if not isinstance(value, list):
        return []
    out: list[str] = []
    for item in value:
        s = _clean_ws(str(item))
        if s:
            out.append(s)
    return out


def _normalize_student_text_for_match(text: str) -> str:
    t = _clean_ws(text).lower()
    # Normalize common unicode punctuation and drop non-word symbols for robust matching.
    t = t.replace("’", "'").replace("‘", "'").replace("“", '"').replace("”", '"').replace("—", "-")
    t = re.sub(r"[^a-z0-9\s]", " ", t)
    t = re.sub(r"\s+", " ", t).strip()
    return t


def _evidence_match_count(evidence_refs: list[str], student_text: str) -> int:
    if not evidence_refs or not student_text:
        return 0
    st = _normalize_student_text_for_match(student_text)
    st_tokens = set(st.split())
    hits = 0
    for ref in evidence_refs:
        token = _normalize_student_text_for_match(ref)
        if len(token) < 10:
            continue
        if token in st:
            hits += 1
            continue
        # Fuzzy fallback: token overlap (handles punctuation/quote variations).
        ref_tokens = [t for t in token.split() if len(t) >= 4]
        if len(ref_tokens) < 3:
            continue
        overlap = sum(1 for t in set(ref_tokens) if t in st_tokens)
        ratio = overlap / max(1, len(set(ref_tokens)))
        if overlap >= 3 and ratio >= 0.5:
            hits += 1
    return hits


def _assignment_requires_workflow_diagram(assignment_text: str) -> bool:
    t = _normalize_student_text_for_match(assignment_text)
    return ("workflow diagram" in t) or ("develop a workflow diagram" in t) or ("draw a workflow" in t)


def _student_has_diagram_evidence(student_text: str, student_chunks: list[dict[str, Any]]) -> bool:
    """
    Return True if the student likely submitted a workflow diagram.
    Checks two signals:
    1. Any image chunk in the describe output (diagram as image/figure in PDF).
    2. Diagram-related keywords in student text (BPMN terms, figure refs, swimlane, etc.).
    """
    # 1. Any image-type chunk present
    for c in student_chunks:
        md = c.get("metadata", {}) or {}
        ctype = str(md.get("content_type", "")).lower()
        if "image" in ctype:
            return True

    # 2. Text keywords indicating a diagram was described or referenced
    t = _normalize_student_text_for_match(student_text)
    diagram_keywords = [
        "swim lane", "swimlane", "swim-lane",
        "workflow diagram", "flow diagram", "process diagram",
        "flowchart", "flow chart",
        "decision point", "decision node",
        "start event", "end event",
        "current state diagram", "as-is diagram", "as is diagram",
        "figure 1", "figure 2", "see figure", "diagram below",
        "attached diagram", "diagram shows", "diagram illustrates",
        "bpmn", "uml diagram",
        "the diagram", "in the diagram", "refer to diagram",
        "step 1", "step 2", "step 3",  # numbered workflow steps
    ]
    return any(kw in t for kw in diagram_keywords)


def _snap_to_grade_band(checklist_pct: float, max_points: float) -> float:
    """Convert a checklist satisfaction percentage to awarded points via the rubric grade band table."""
    BANDS = [
        (95, 1.000),
        (90, 0.933),
        (85, 0.900),
        (80, 0.833),
        (75, 0.800),
        (70, 0.733),
        (65, 0.700),
        (60, 0.633),
        (55, 0.600),
        (0,  0.500),
    ]
    for threshold, multiplier in BANDS:
        if checklist_pct >= threshold:
            return round(max_points * multiplier * 2) / 2  # round to nearest 0.5
    return round(max_points * 0.500 * 2) / 2


def normalize_grade_result(
    result: dict[str, Any],
    *,
    rubric_criteria: list[dict[str, Any]],
    expected_sections: list[str],
    student_file: str,
    student_text: str,
    student_chunks: list[dict[str, Any]],
    assignment_text: str,
) -> dict[str, Any]:
    if not isinstance(result, dict):
        return {
            "student_file": student_file,
            "overall_score": 0.0,
            "overall_feedback": "Invalid grading response format.",
            "questions": [],
            "parse_error": True,
        }

    # If model did not return criterion_scores, fallback to older schema.
    criterion_scores_raw = result.get("criterion_scores")
    if not isinstance(criterion_scores_raw, list) or not criterion_scores_raw:
        criterion_scores_raw = result.get("criterionScores")
    if not isinstance(criterion_scores_raw, list) or not criterion_scores_raw:
        criterion_scores_raw = result.get("criteria_scores")
    if not isinstance(criterion_scores_raw, list) or not criterion_scores_raw:
        criterion_scores_raw = result.get("criteria")
    if isinstance(criterion_scores_raw, dict):
        converted: list[dict[str, Any]] = []
        for key, value in criterion_scores_raw.items():
            if isinstance(value, dict):
                item = dict(value)
                if "criterion_id" not in item and "criterionId" not in item:
                    item["criterion_id"] = str(key)
                converted.append(item)
            else:
                score_val = _to_float(value)
                if score_val is not None:
                    converted.append({"criterion_id": str(key), "awarded_points": score_val})
        criterion_scores_raw = converted
    if (not isinstance(criterion_scores_raw, list) or not criterion_scores_raw) and _is_single_criterion_fragment(result):
        criterion_scores_raw = [result]
    if not isinstance(criterion_scores_raw, list) or not criterion_scores_raw:
        criterion_scores_raw = []
    if not isinstance(criterion_scores_raw, list) or not criterion_scores_raw:
        score_breakdown_raw = result.get("score_breakdown", result.get("scoreBreakdown", {}))
        score_breakdown: dict[str, float] = {}
        if isinstance(score_breakdown_raw, dict):
            for key, raw_val in score_breakdown_raw.items():
                val = _to_float(raw_val)
                if val is not None:
                    score_breakdown[str(key)] = round(max(0.0, val), 2)
        overall_score = round(min(100.0, max(0.0, sum(score_breakdown.values()))), 2) if score_breakdown else 0.0
        return {
            "student_file": str(result.get("student_file", student_file)).strip() or student_file,
            "overall_score": overall_score,
            "overall_feedback": str(
                result.get("overall_feedback", result.get("overallFeedback", "Automated grading completed."))
            ).strip(),
            "score_breakdown": score_breakdown,
            "criterion_details": [],
            "section_coverage": [],
            "strengths": _to_string_list(result.get("strengths", result.get("key_strengths", []))),
            "gaps": _to_string_list(result.get("gaps", result.get("key_gaps", []))),
            "action_items": _to_string_list(result.get("action_items", result.get("actionItems", []))),
            "confidence": _to_float(result.get("confidence")) if result.get("confidence") is not None else None,
            "policy_caps_applied": [],
            "questions": [],
            "parse_error": bool(result.get("parse_error")),
        }

    criteria_by_id = {str(c.get("criterion_id")): c for c in rubric_criteria}
    criteria_by_name = {_slugify(str(c.get("criterion_name", ""))): c for c in rubric_criteria}

    criterion_details: list[dict[str, Any]] = []
    score_breakdown: dict[str, float] = {}

    # Build a lookup of model entries by criterion_id or criterion_name slug.
    model_entries: dict[str, dict[str, Any]] = {}
    for item in criterion_scores_raw:
        if not isinstance(item, dict):
            continue
        cid = str(item.get("criterion_id", item.get("criterionId", ""))).strip()
        cname = _slugify(str(item.get("criterion_name", item.get("criterionName", item.get("criterion", "")))))
        if cid:
            model_entries[cid] = item
        elif cname:
            model_entries[cname] = item

    for crit in rubric_criteria:
        cid = str(crit.get("criterion_id", "")).strip()
        cname = str(crit.get("criterion_name", "")).strip()
        max_points = float(crit.get("max_points", 0.0) or 0.0)

        model_item = model_entries.get(cid) or model_entries.get(_slugify(cname), {})
        raw_awarded = _to_float(
            model_item.get(
                "awarded_points",
                model_item.get("awardedPoints", model_item.get("score", model_item.get("points"))),
            )
        )
        awarded = 0.0 if raw_awarded is None else max(0.0, min(max_points, raw_awarded))

        # Grade band snapping: if model returned checklist_pct, override awarded_points
        # with the deterministic band value so grading matches the professor's rubric bands.
        checklist_pct = _to_float(model_item.get("checklist_pct"))
        grade_band_snapped = False
        if checklist_pct is not None:
            band_pts = _snap_to_grade_band(float(max(0.0, min(100.0, checklist_pct))), max_points)
            awarded = band_pts
            grade_band_snapped = True

        evidence_refs = _to_string_list(
            model_item.get("evidence_refs", model_item.get("evidenceRefs", model_item.get("evidence", [])))
        )
        missing_items = _to_string_list(
            model_item.get("missing_items", model_item.get("missingItems", model_item.get("gaps", [])))
        )
        justification = _clean_ws(str(model_item.get("justification", model_item.get("rationale", ""))))
        evidence_hits = _evidence_match_count(evidence_refs, student_text)

        # Anti-hallucination guard: cap criterion at 60% of max when no evidence found.
        # Skip this cap for image-heavy submissions — evidence may be in diagrams/tables, not text.
        has_image_chunks = any(
            "image" in str(c.get("metadata", {}).get("content_type", "")).lower()
            for c in student_chunks
        )
        no_evidence_cap_applied = False
        if evidence_hits == 0 and not has_image_chunks and not grade_band_snapped:
            cap_val = round(max_points * 0.60, 2)
            if awarded > cap_val:
                awarded = cap_val
                no_evidence_cap_applied = True

        key = cname
        score_breakdown[key] = round(awarded, 2)
        criterion_details.append(
            {
                "criterion_id": cid,
                "criterion_name": cname,
                "max_points": round(max_points, 2),
                "awarded_points": round(awarded, 2),
                "checklist_pct": round(float(checklist_pct), 1) if checklist_pct is not None else None,
                "grade_band_snapped": grade_band_snapped,
                "justification": justification,
                "evidence_refs": evidence_refs,
                "evidence_match_count": evidence_hits,
                "missing_items": missing_items,
                "no_evidence_cap_applied": no_evidence_cap_applied,
            }
        )

    # Section coverage (deterministic structure)
    section_coverage_raw = result.get("section_coverage", result.get("sectionCoverage", []))
    section_coverage: list[dict[str, Any]] = []
    if isinstance(section_coverage_raw, list):
        for item in section_coverage_raw:
            if not isinstance(item, dict):
                continue
            sid = _clean_ws(str(item.get("section_id", item.get("sectionId", ""))))
            status = _clean_ws(str(item.get("status", "")).lower())
            if status not in {"addressed", "partial", "missing"}:
                status = "partial" if status else ""
            evidence_refs = _to_string_list(item.get("evidence_refs", item.get("evidenceRefs", [])))
            if sid and status:
                section_coverage.append(
                    {"section_id": sid, "status": status, "evidence_refs": evidence_refs}
                )

    # Deterministic scoring
    raw_total = round(sum(score_breakdown.values()), 2)
    final_score = raw_total
    policy_caps: list[str] = []

    # Cap for missing required workflow diagram deliverable.
    if _assignment_requires_workflow_diagram(assignment_text):
        if not _student_has_diagram_evidence(student_text, student_chunks):
            if final_score > 78.0:
                final_score = 78.0
                policy_caps.append("missing_required_workflow_diagram_cap_78")

    # Cap for missing assignment sections.
    if expected_sections:
        sec_map = {str(s.get("section_id", "")).upper(): str(s.get("status", "")) for s in section_coverage}
        missing_count = 0
        partial_count = 0
        for s in expected_sections:
            st = sec_map.get(s.upper(), "missing")
            if st == "missing":
                missing_count += 1
            elif st == "partial":
                partial_count += 1
        if missing_count > 0 or partial_count > 1:
            # deterministic conservative cap
            section_cap = max(55.0, 100.0 - (missing_count * 15.0) - (partial_count * 6.0))
            if final_score > section_cap:
                final_score = round(section_cap, 2)
                policy_caps.append(
                    f"section_coverage_cap_missing_{missing_count}_partial_{partial_count}"
                )

    overall_feedback = _clean_ws(str(result.get("overall_feedback", result.get("overallFeedback", ""))))
    if not overall_feedback:
        overall_feedback = "Automated grading completed."

    strengths = _to_string_list(result.get("strengths", result.get("key_strengths", [])))
    gaps = _to_string_list(result.get("gaps", result.get("key_gaps", [])))
    action_items = _to_string_list(result.get("action_items", result.get("actionItems", [])))

    return {
        "student_file": str(result.get("student_file", student_file)).strip() or student_file,
        "overall_score": round(max(0.0, min(100.0, final_score)), 2),
        "overall_feedback": overall_feedback,
        "score_breakdown": score_breakdown,
        "criterion_details": criterion_details,
        "section_coverage": section_coverage,
        "strengths": strengths,
        "gaps": gaps,
        "action_items": action_items,
        "confidence": _to_float(result.get("confidence")) if result.get("confidence") is not None else None,
        "policy_caps_applied": policy_caps,
        "questions": [],
        "parse_error": bool(result.get("parse_error")),
    }


# ---------------------------------------------------------------------------
# CLI
# ---------------------------------------------------------------------------

def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(description="Grade a student submission via OpenAI")
    parser.add_argument("--retrieval-jsonl", required=True, help="Path to retrieval_results.jsonl")
    parser.add_argument("--chunks-jsonl", required=True, help="Path to chunks.jsonl")
    parser.add_argument(
        "--student-path",
        default=None,
        help="Substring to filter to a single student file (e.g. 'Student 1')",
    )
    parser.add_argument("--out-dir", default=None, help="Output directory (default: alongside retrieval-jsonl)")
    parser.add_argument(
        "--grading-provider",
        default="openai",
        choices=list(GRADING_PROVIDERS.keys()),
        help="LLM provider for grading (openai, gemini, anthropic)",
    )
    parser.add_argument("--model", default=None, help="Model name (default: provider's default model)")
    parser.add_argument("--max-lecture-chars", type=int, default=12000)
    parser.add_argument("--max-student-chars", type=int, default=20000)
    parser.add_argument(
        "--rubric-file",
        default=None,
        help="Optional rubric text file used to anchor scoring.",
    )
    parser.add_argument(
        "--assignment-file",
        default=None,
        help="Text file with the actual assignment instructions/questions.",
    )
    return parser.parse_args()


def run_grading(
    *,
    retrieval_jsonl: Path,
    chunks_jsonl: Path,
    student_path_filter: str | None,
    out_dir: Path,
    model: str,
    grading_provider: str = "openai",
    max_lecture_chars: int,
    max_student_chars: int,
    rubric_file: Path | None = None,
    assignment_file: Path | None = None,
) -> Path:
    if grading_provider not in GRADING_PROVIDERS:
        raise RuntimeError(f"Unknown grading provider: {grading_provider}. Use: {list(GRADING_PROVIDERS.keys())}")

    key_name, call_fn = GRADING_PROVIDERS[grading_provider]
    api_key = get_api_key(key_name)
    if not api_key:
        raise RuntimeError(f"{key_name.upper()}_API_KEY not set. Add it to your .env file.")

    if not retrieval_jsonl.exists():
        raise RuntimeError(f"retrieval_results.jsonl not found: {retrieval_jsonl}")
    if not chunks_jsonl.exists():
        raise RuntimeError(f"chunks.jsonl not found: {chunks_jsonl}")

    print(f"Loading student content from: {chunks_jsonl}")
    student_text, student_chunks = load_student_content(chunks_jsonl, student_path_filter)
    if not student_text:
        raise RuntimeError("No student content found. Check --student-path filter or chunks.jsonl.")

    # Determine the student filename for the report.
    student_file = "unknown"
    if student_chunks:
        student_file = student_chunks[0].get("metadata", {}).get("filename", "unknown")
    student_file_for_model = anonymize_label(student_file)

    print(f"Student file: {student_file}  |  {len(student_text):,} chars  |  {len(student_chunks)} chunks")

    print(f"Loading lecture context from: {retrieval_jsonl}")
    lecture_context = load_lecture_context(retrieval_jsonl, student_path_filter, max_lecture_chars)
    print(f"Lecture context: {len(lecture_context):,} chars")

    if not lecture_context:
        print("WARNING: No lecture context found. Grading without lecture reference.")
        lecture_context = "(No lecture context available)"

    rubric_text = read_text_file(rubric_file)
    if rubric_file and not rubric_text:
        print(f"WARNING: rubric file not found/empty: {rubric_file}")
    rubric_criteria: list[dict[str, Any]] = []
    if rubric_file and rubric_file.suffix.lower() == ".json":
        # Generated rubric JSON — parse directly so criteria names and points are exact.
        try:
            rubric_data = json.loads(rubric_file.read_text(encoding="utf-8"))
            raw_criteria = rubric_data.get("criteria", [])
            if raw_criteria:
                rubric_criteria = [
                    {
                        "criterion_id": f"C{i + 1}",
                        "criterion_name": c.get("criterion_name", f"Criterion {i + 1}"),
                        "max_points": float(c.get("max_points", 0)),
                        "checklist_items": c.get("checklist_items", []),
                    }
                    for i, c in enumerate(raw_criteria)
                ]
                print(f"JSON rubric loaded: {len(rubric_criteria)} criteria, "
                      f"{sum(c['max_points'] for c in rubric_criteria):.0f} total pts")
        except Exception as exc:
            print(f"WARNING: Failed to parse JSON rubric: {exc}")
    elif rubric_file and rubric_file.suffix.lower() == ".docx":
        rubric_criteria = extract_rubric_criteria_from_docx(rubric_file)
    if not rubric_criteria:
        rubric_criteria = extract_rubric_criteria(rubric_text)
    # AI fallback — when rubric is prose/free-text that regex parsers cannot handle.
    # The parser provider is chosen independently of the grading provider because
    # ai_extract_rubric_criteria only supports anthropic/openai (not gemini).
    if not rubric_criteria and rubric_text:
        if grading_provider in ("anthropic", "openai"):
            parser_provider = grading_provider
            parser_api_key = api_key
        elif get_api_key("anthropic"):
            parser_provider = "anthropic"
            parser_api_key = get_api_key("anthropic")
        elif get_api_key("openai"):
            parser_provider = "openai"
            parser_api_key = get_api_key("openai")
        else:
            parser_provider = None
            parser_api_key = None

        if parser_provider:
            print(
                f"INFO: Regex rubric parsers found no criteria — trying AI rubric "
                f"extraction via {parser_provider}..."
            )
            rubric_criteria = ai_extract_rubric_criteria(
                rubric_text, parser_api_key, parser_provider,
            )
        else:
            print(
                "INFO: Regex rubric parsers found no criteria, and AI rubric extraction "
                "is unavailable (set ANTHROPIC_API_KEY or OPENAI_API_KEY to enable). "
                "Falling back to default criteria."
            )
    _rubric_from_file = bool(rubric_criteria)  # True if successfully parsed from actual rubric
    if not rubric_criteria:
        rubric_criteria = list(DEFAULT_RUBRIC_CRITERIA)
        print("WARNING: No rubric file provided or criteria could not be parsed. "
              "Using generic default criteria — upload a rubric DOCX/PDF for accurate assignment-specific grading.")
    else:
        print(
            "Rubric criteria parsed: "
            + ", ".join(
                [f"{c['criterion_id']}={c['criterion_name']} ({int(c['max_points'])})" for c in rubric_criteria]
            )
        )

    assignment_text = read_text_file(assignment_file)
    if assignment_file and not assignment_text:
        print(f"WARNING: assignment file not found/empty: {assignment_file}")
    if assignment_text:
        print(f"Assignment instructions loaded: {len(assignment_text):,} chars")
    else:
        print("WARNING: No assignment instructions — model will infer questions from student text.")

    # Prefer sections from assignment instructions (authoritative) over student text.
    # This ensures that even if a student's text has no numbered headings, the
    # grader uses the actual assignment questions (Q1-Q4) as the grading template.
    # Sections derived from the assignment are stable across students, so they are
    # safe to embed in the cached system block. Sections derived from a student's
    # own text vary per student and would invalidate the cache on every call —
    # those are kept out of the cached block and used only for downstream coverage.
    if assignment_text:
        assignment_sections = detect_expected_sections(assignment_text)
        if assignment_sections:
            print(f"Expected sections (from assignment): {assignment_sections}")
        else:
            print("WARNING: Could not detect section IDs from assignment file.")
    else:
        assignment_sections = []

    expected_sections = assignment_sections
    if not expected_sections:
        expected_sections = detect_expected_sections(student_text)
        if expected_sections:
            print(f"Expected sections (from student text — not cached): {expected_sections}")

    system_blocks = build_system_blocks(
        rubric_text=rubric_text,
        rubric_criteria=rubric_criteria,
        assignment_text=assignment_text,
        expected_sections=assignment_sections,
        enable_cache=(grading_provider == "anthropic"),
    )
    user_msg = build_user_message(
        student_text=student_text,
        lecture_context=lecture_context,
        student_file=student_file_for_model,
        max_student_chars=max_student_chars,
    )

    print(f"Calling {grading_provider} ({model}) ...")
    response = call_fn(
        model=model,
        api_key=api_key,
        system=system_blocks,
        user=user_msg,
    )

    usage = response["usage"]
    cache_read = usage.get("cache_read_input_tokens", 0)
    cache_write = usage.get("cache_creation_input_tokens", 0)
    cache_info = f"  cache_write: {cache_write:,}  cache_read: {cache_read:,}" if (cache_read or cache_write) else ""
    print(f"Tokens used — input: {usage['input_tokens']:,}  output: {usage['output_tokens']:,}{cache_info}")

    normalized = normalize_grade_result(
        response["result"],
        rubric_criteria=rubric_criteria,
        expected_sections=expected_sections,
        student_file=student_file,
        student_text=student_text,
        student_chunks=student_chunks,
        assignment_text=assignment_text,
    )

    output = {
        "grading_model": response["model"],
        "student_path_filter": student_path_filter,
        "chunks_jsonl": str(chunks_jsonl),
        "retrieval_jsonl": str(retrieval_jsonl),
        "rubric_file": str(rubric_file) if rubric_file else None,
        "rubric_criteria": rubric_criteria,
        "rubric_used_generic_defaults": not _rubric_from_file,
        "assignment_file": str(assignment_file) if assignment_file else None,
        "expected_sections": expected_sections,
        "token_usage": usage,
        **normalized,
    }
    raw_result = response.get("result", {})
    if isinstance(raw_result, dict):
        output["llm_result_keys"] = sorted([str(k) for k in raw_result.keys()])
        if not normalized.get("criterion_details"):
            output["llm_unmapped_result"] = raw_result
    if isinstance(raw_result, dict) and raw_result.get("parse_error"):
        output["llm_parse_error"] = True
        raw_text = str(raw_result.get("raw_response", "")).strip()
        if raw_text:
            output["llm_raw_response"] = raw_text[:100000]

    out_path = out_dir / "grades.json"
    write_json(out_path, output)
    print(f"\nGrades written to: {out_path}")

    # Print a quick summary.
    result = normalized
    if not result.get("parse_error"):
        print(f"\nOverall score: {result.get('overall_score', 'N/A')} / 100")
        print(f"Overall feedback: {result.get('overall_feedback', '')}")
        breakdown = result.get("score_breakdown", {})
        if isinstance(breakdown, dict) and breakdown:
            print("\nScore breakdown:")
            for k, v in breakdown.items():
                print(f"  {k}: {v}")
        caps = result.get("policy_caps_applied", [])
        if isinstance(caps, list) and caps:
            print("\nPolicy caps applied:")
            for cap in caps:
                print(f"  - {cap}")
        details = result.get("criterion_details", [])
        if isinstance(details, list) and details:
            print(f"\nCriterion details ({len(details)}):")
            for d in details:
                print(
                    f"  [{d.get('criterion_id')}] {d.get('criterion_name')}: "
                    f"{d.get('awarded_points')}/{d.get('max_points')} "
                    f"(evidence matches={d.get('evidence_match_count')})"
                )
    else:
        print(f"\nWARNING: Could not parse JSON from {grading_provider} response. Raw output saved in grades.json.")

    return out_path


def main() -> int:
    load_environment()
    args = parse_args()

    retrieval_jsonl = Path(args.retrieval_jsonl).expanduser().resolve()
    chunks_jsonl = Path(args.chunks_jsonl).expanduser().resolve()
    out_dir = Path(args.out_dir).expanduser().resolve() if args.out_dir else retrieval_jsonl.parent
    rubric_file = Path(args.rubric_file).expanduser().resolve() if args.rubric_file else None
    assignment_file = Path(args.assignment_file).expanduser().resolve() if args.assignment_file else None

    grading_provider = args.grading_provider
    model = args.model or DEFAULT_GRADING_MODELS.get(grading_provider, "gpt-4o-2024-11-20")

    run_grading(
        retrieval_jsonl=retrieval_jsonl,
        chunks_jsonl=chunks_jsonl,
        student_path_filter=args.student_path,
        out_dir=out_dir,
        model=model,
        grading_provider=grading_provider,
        max_lecture_chars=int(args.max_lecture_chars),
        max_student_chars=int(args.max_student_chars),
        rubric_file=rubric_file,
        assignment_file=assignment_file,
    )
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
